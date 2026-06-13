from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('/Users/alesso/dev/university/PDC/PESTEL_Ingenieria_Economica_Argentina.docx')

BLUE = '1F4E79'
DARK = '0B2545'
LIGHT = 'EAF2F8'
GRAY = 'F2F4F7'
MUTED = '666666'
WHITE = 'FFFFFF'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text_color(cell, color):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for m, v in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_table_borders(table, color='D9E2EC'):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '6')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_meta(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(label + ': ')
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run(value)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color='B7C9D9')
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p2 = cell.add_paragraph(text)
    p2.paragraph_format.space_after = Pt(0)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.10
for name in ['List Bullet', 'List Number']:
    styles[name].font.name = 'Calibri'
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(4)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Análisis PESTEL Integrado')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Ingeniería Económica - Argentina 2026')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_meta(doc, 'Tema integrador', 'PBI, comercio internacional, flujo de inversiones, salarios, pobreza, innovación y marco institucional argentino.')
add_meta(doc, 'Base documental', 'Presentaciones e informes de Ingeniería Económica provistos por el alumno en Desktop.')
add_meta(doc, 'Enfoque', 'Diagnóstico estratégico del entorno macroeconómico argentino usando la matriz PESTEL.')

add_callout(doc, 'Idea fuerza', 'Argentina posee recursos productivos, capital humano y sectores exportadores con potencial, pero su crecimiento queda condicionado por inflación, restricción externa, baja inversión, informalidad, desigualdad social y reglas institucionales inestables.')

add_h1(doc, 'Matriz PESTEL')
intro = doc.add_paragraph('La matriz sintetiza los factores externos que condicionan el desempeño económico argentino y la toma de decisiones de inversión. Integra los ejes trabajados en los materiales: PBI, comercio internacional, inversión, salarios, pobreza, propiedad intelectual y economía del conocimiento.')
intro.paragraph_format.space_after = Pt(8)

data = [
    ('Político', 'Estabilidad macro, política fiscal, credibilidad institucional, programas sociales y orientación del Estado.', 'El orden fiscal y monetario puede mejorar expectativas, pero si el ajuste reduce inversión pública o protección social puede limitar actividad y bienestar.', 'Riesgo país, conflictividad distributiva, continuidad de reglas, sesgo del gasto y capacidad estatal.'),
    ('Económico', 'Inflación, PBI real, inversión, balanza de pagos, salarios reales, crédito y restricción externa.', 'La economía puede rebotar, pero necesita inversión y divisas estables. El superávit comercial no alcanza si servicios, rentas e importaciones presionan la cuenta corriente.', 'Inflación persistente, falta de reservas, baja inversión, volatilidad cambiaria y caída del salario real.'),
    ('Social', 'Pobreza, informalidad, desigualdad regional, capital humano, empleo formal y salarios.', 'El crecimiento del PBI no garantiza bienestar si convive con pobreza infantil, empleo informal o pérdida de poder adquisitivo.', 'Trabajador pobre, informalidad superior al 40%, baja cobertura social, pobreza infantil y brecha educativa.'),
    ('Tecnológico', 'Software, servicios basados en conocimiento, IA, I+D, infraestructura digital y propiedad intelectual.', 'La economía del conocimiento puede generar divisas y empleo calificado, pero la brecha digital limita el derrame hacia sectores vulnerables.', 'Fuga de talento, baja escala de I+D, rezago 5G, ciberseguridad, patentes y regulación de IA.'),
    ('Ecológico', 'Sequías, clima, agro, energía, minería, litio, ESG, data centers y uso de recursos.', 'El ambiente afecta directamente exportaciones, disponibilidad de divisas y PBI. También abre oportunidades en energía, minería, AgTech y transición energética.', 'Sequías, estrés hídrico, consumo energético de IA/data centers, conflictos socioambientales y exigencias ESG.'),
    ('Legal', 'Seguridad jurídica, régimen tributario, reglas laborales, propiedad intelectual, RIGI/RIMI y regulaciones cambiarias.', 'Las reglas definen costos, incentivos y horizontes de inversión. Regímenes promocionales ayudan, pero no reemplazan estabilidad macro e institucional.', 'Presión tributaria, litigiosidad laboral, cambios normativos, demoras del INPI y complejidad administrativa.'),
]

table = doc.add_table(rows=1, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
headers = ['Factor', 'Aspectos clave', 'Impacto económico', 'Riesgos principales']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_shading(cell, BLUE)
    set_cell_text_color(cell, WHITE)
    set_cell_margins(cell)
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_repeat_table_header(table.rows[0])
for row in data:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i == 0:
            set_cell_shading(cells[i], GRAY)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor.from_string(DARK)
set_table_borders(table)
set_col_widths(table, [1.0, 1.85, 2.0, 1.65])

doc.add_paragraph()
add_h1(doc, 'Desarrollo por factor')

sections = [
    ('Político', [
        'El entorno político incide sobre inversión, consumo y financiamiento porque define expectativas sobre política fiscal, monetaria y cambiaria.',
        'Los materiales de PBI remarcan que la estabilidad puede crear condiciones para crecimiento sostenible, pero el ajuste fiscal concentrado en gasto e inversión pública puede enfriar la actividad en el corto plazo.',
        'En pobreza, el Estado aparece como actor necesario para corregir fallas de mercado: transferencias como AUH y Tarjeta Alimentar reducen indigencia, aunque persiste un sesgo del gasto hacia adultos mayores frente a la pobreza infantil.',
        'En comercio internacional, la restricción externa exige políticas de largo plazo: no alcanza con superávit de bienes si la cuenta corriente se deteriora por servicios, turismo emisivo y renta de inversión.'
    ]),
    ('Económico', [
        'El PBI es útil para medir producción, pero no alcanza para evaluar bienestar: no muestra distribución del ingreso, informalidad, trabajo doméstico no remunerado ni costos ambientales.',
        'La economía argentina muestra volatilidad histórica: expansión, crisis, rebotes y nuevas caídas. En 2025 se observa recuperación parcial del PBI real cercana al 4,4%, pero heterogénea entre sectores.',
        'El sector externo es una restricción central: marzo de 2026 muestra exportaciones por USD 8.645 millones y saldo comercial positivo, pero el informe de comercio advierte que la cuenta corriente puede seguir frágil por servicios y rentas.',
        'La inflación reduce salario real, encarece la canasta básica y distorsiona decisiones de inversión. Los materiales de PBI mencionan una inflación proyectada por el FMI de 30,4% anual para 2026 y un IPC de marzo de 2026 de 3,4%.',
        'En inversión, el uso de VAN, TIR, flujo de fondos, costo de oportunidad y tasa de descuento resulta clave para decidir proyectos en contextos de riesgo, inflación y volatilidad.'
    ]),
    ('Social', [
        'La pobreza se presenta como fenómeno multidimensional: ingresos insuficientes, falta de activos, déficit habitacional, inseguridad alimentaria, baja cobertura de salud y vulnerabilidad ante shocks.',
        'El material sobre pobreza destaca una tasa oficial de pobreza de 28,2% e indigencia de 6,3% para el segundo semestre de 2025, pero con pobreza infantil urbana de 53,6%, lo que compromete el capital humano futuro.',
        'El mercado laboral muestra una brecha marcada: pocos empleos privados formales sostienen una estructura amplia, mientras la informalidad limita productividad, protección social y recaudación.',
        'El informe de salarios IT muestra que sectores de alta productividad pueden pagar remuneraciones muy superiores al promedio, especialmente cuando exportan servicios; el contraste con trabajadores informales refuerza la segmentación social.',
        'La recuperación del PBI solo se transforma en bienestar si mejora el empleo formal, el salario real, la educación, la salud y el acceso a infraestructura básica.'
    ]),
    ('Tecnológico', [
        'La tecnología aparece como oportunidad de productividad y exportación. El sector SSI transforma talento en divisas y puede aliviar la restricción externa mediante servicios de alto valor agregado.',
        'Los materiales de comercio mencionan exportaciones SSI por USD 2.449 millones en 2023, crecimiento anual relevante entre 2015 y 2022 y un rol potencial en diversificación exportadora.',
        'La IA, fintech, salud, AgTech y e-commerce son verticales dinámicas, pero su impacto puede quedar concentrado si no se resuelven conectividad, educación digital, infraestructura y financiamiento.',
        'La brecha digital profundiza desigualdad: en hogares pobres, el acceso limitado a internet de calidad y computadoras reduce habilidades futuras y empleabilidad.',
        'La propiedad intelectual es crítica: la Ley de Economía del Conocimiento y los regímenes de patentes/marcas incentivan innovación, aunque los tiempos administrativos y la discusión sobre IA/autoria plantean desafíos.'
    ]),
    ('Ecológico / ambiental', [
        'El ambiente no es externo a la economía: sequías, clima, agua y energía afectan producción agropecuaria, exportaciones, reservas, inflación de alimentos y recaudación.',
        'El PBI no descuenta degradación ambiental ni agotamiento de recursos, por eso puede sobreestimar bienestar si el crecimiento se basa en deterioro ambiental.',
        'Argentina tiene oportunidades en litio, minería, Vaca Muerta, energías renovables, AgTech y soluciones ESG, pero requieren infraestructura, licencia social y reglas ambientales claras.',
        'La expansión de IA y data centers eleva demanda eléctrica y de agua para enfriamiento; esto condiciona localización y sostenibilidad de inversiones tecnológicas.',
        'Los hogares pobres son más vulnerables a shocks climáticos porque suelen carecer de seguros, ahorro, infraestructura y vivienda segura.'
    ]),
    ('Legal / institucional', [
        'La seguridad jurídica, la estabilidad tributaria, las reglas laborales y las regulaciones cambiarias determinan costos de formalización e inversión.',
        'Los materiales de inversión mencionan incentivos como RIGI para grandes inversiones y RIMI para PyMEs, con beneficios fiscales, estabilidad y mecanismos de recuperación de IVA/amortización.',
        'La Ley de Economía del Conocimiento ofrece beneficios fiscales a software, biotecnología y actividades intensivas en conocimiento, a cambio de I+D, empleo y exportaciones.',
        'El marco laboral incide sobre contratación formal: reformas como fondos de cese, períodos de prueba y reducción de contribuciones pueden ayudar, pero no resuelven por sí solas déficit de cualificaciones.',
        'En propiedad intelectual, Argentina cuenta con Ley 11.723 y Ley 24.481, pero enfrenta desafíos de modernización, tiempos del INPI, enforcement y adecuación a tratados internacionales.'
    ]),
]

for title, bullets in sections:
    add_h2(doc, title)
    for item in bullets:
        add_bullet(doc, item)

add_h1(doc, 'Síntesis estratégica')
add_callout(doc, 'Diagnóstico integrado', 'El problema central no es la ausencia de sectores con potencial, sino la dificultad de convertirlos en crecimiento sostenido e inclusivo. Argentina combina agroindustria, energía, minería, servicios profesionales, software y capital humano, pero enfrenta inflación, restricción externa, baja inversión, informalidad, brechas sociales y reglas inestables.')

for item in [
    'Oportunidad principal: diversificar exportaciones hacia energía, minería, agroindustria sofisticada y servicios basados en conocimiento para generar divisas menos dependientes del ciclo de commodities.',
    'Riesgo principal: que el crecimiento sea solo un rebote estadístico del PBI y no se traduzca en empleo formal, productividad, salario real y reducción de pobreza.',
    'Condición necesaria: estabilización macroeconómica, pero acompañada por inversión, infraestructura, educación, capacitación tecnológica y reglas previsibles.',
    'Condición social: proteger primera infancia y capital humano para evitar que pobreza, brecha digital y mala nutrición limiten la productividad futura.',
]:
    add_bullet(doc, item)

add_h1(doc, 'Fuentes usadas')
source_rows = [
    ('Comercio Internacional', 'Comercio-Internacional-Presentacion (1).pptx; IE - Comercio Internacional Informe.pdf'),
    ('Flujo de inversiones', 'Flujo_Inversiones_Presentacion_Economia2026_V4.pptx'),
    ('PBI', 'IE - TPI - PBI.pdf; IE- Presentacion - PBI.pdf'),
    ('Pobreza', 'PDF - IE - TP - Pobreza.pdf; index.html'),
    ('Salarios y mercado IT', 'informe_salarios_argentina.pdf'),
    ('Propiedad intelectual / marco institucional', 'INFORME IE.pdf; PP IE.pptx'),
]
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
t.rows[0].cells[0].text = 'Eje'
t.rows[0].cells[1].text = 'Archivos considerados'
for cell in t.rows[0].cells:
    set_cell_shading(cell, BLUE)
    set_cell_text_color(cell, WHITE)
    set_cell_margins(cell)
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for eje, archivo in source_rows:
    cells = t.add_row().cells
    cells[0].text = eje
    cells[1].text = archivo
    set_cell_margins(cells[0])
    set_cell_margins(cells[1])
set_table_borders(t)
set_col_widths(t, [1.8, 4.7])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
r = p.add_run('Nota: documento elaborado a partir de los archivos provistos localmente. Los datos numéricos se mantienen como aparecen en los materiales de la materia.')
r.italic = True
r.font.color.rgb = RGBColor.from_string(MUTED)

# Footer
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Ingeniería Económica - PESTEL integrado')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor.from_string(MUTED)

# Core properties
doc.core_properties.title = 'Análisis PESTEL Integrado - Ingeniería Económica'
doc.core_properties.subject = 'Argentina 2026'
doc.core_properties.author = 'Codex'

doc.save(OUT)
print(OUT)
