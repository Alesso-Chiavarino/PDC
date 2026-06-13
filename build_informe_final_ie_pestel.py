from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/alesso/dev/university/PDC/Informe_Final_IE_PESTEL_Argentina_SSI.docx')

NAVY = '0B2545'
BLUE = '1F4D78'
LIGHT_BLUE = 'EAF3FA'
LIGHT_GRAY = 'F2F4F7'
PALE_GREEN = 'E7F3EA'
PALE_GOLD = 'FFF4D6'
PALE_RED = 'FBE7E7'
WHITE = 'FFFFFF'
MUTED = '555555'
BORDER = 'D7DEE6'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, val in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn(f'w:{side}'))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = tbl_borders.find(qn(f'w:{edge}'))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tbl_borders.append(node)
        node.set(qn('w:val'), 'single')
        node.set(qn('w:sz'), '6')
        node.set(qn('w:space'), '0')
        node.set(qn('w:color'), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement('w:tblHeader')
    tbl_header.set(qn('w:val'), 'true')
    tr_pr.append(tbl_header)


def set_column_widths(table, widths_in):
    for row in table.rows:
        for idx, width in enumerate(widths_in):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def style_run(run, bold=False, color=None, size=None, italic=False):
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(caption)
        style_run(r, bold=True, color=NAVY, size=9.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = str(h)
        set_cell_shading(hdr_cells[i], BLUE)
        set_cell_margins(hdr_cells[i])
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                style_run(r, bold=True, color=WHITE, size=9.2)
    repeat_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                set_cell_shading(cells[i], LIGHT_GRAY)
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        style_run(r, bold=True, color=NAVY, size=9.2)
            else:
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.2)
    set_table_borders(table)
    set_column_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Argentina como destino de inversion')
    style_run(r, bold=True, color=NAVY, size=24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run('Informe tecnico PESTEL del entorno empresarial con foco aplicado en SSI')
    style_run(r, italic=True, color=MUTED, size=13)

    add_callout(doc, 'Nota metodologica', 'El informe fue elaborado exclusivamente con los documentos, informes, PDF y presentaciones entregados para la materia Ingenieria Economica. No se incorporan busquedas web ni datos externos nuevos. Por indicacion de trabajo, se excluye la dimension ambiental y se concentra el analisis en factores politicos, economicos, sociales, tecnologicos y legales.')

    for label, text in [
        ('Objeto', 'evaluar el estado actual de Argentina como destino de inversion empresarial.'),
        ('Caso testigo', 'Software y Servicios Informaticos (SSI), por su capacidad de transformar talento en divisas y representar una actividad dinamica de la economia del conocimiento.'),
        ('Criterio de lectura', 'identificar oportunidades, riesgos e impactos sobre decisiones de inversion, financiamiento, localizacion, contratacion y recupero del capital.'),
    ]:
        add_bullet(doc, f'{label}: {text}')


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        style_run(r, bold=True, color=BLUE, size=16)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        style_run(r, bold=True, color=NAVY, size=13)
    return p


def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        style_run(r, bold=True, color='434343', size=12)
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=190, bottom=150, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, bold=True, color=NAVY, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    p2.add_run(body)
    set_table_borders(table, 'B7C9D8')
    set_column_widths(table, [6.3])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def page_break(doc):
    doc.add_page_break()


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    styles['Normal'].font.size = Pt(10.8)
    styles['Normal'].paragraph_format.space_after = Pt(7)
    styles['Normal'].paragraph_format.line_spacing = 1.10
    for style_name in ['List Bullet', 'List Number']:
        styles[style_name].font.name = 'Calibri'
        styles[style_name].font.size = Pt(10.8)
        styles[style_name].paragraph_format.space_after = Pt(4)
        styles[style_name].paragraph_format.line_spacing = 1.10
    return doc


quant_rows = [
    ('Cuenta corriente 1T 2025', '-USD 5.191 millones', 'El saldo externo es fragil aunque existan exportaciones de bienes.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Saldo de bienes 1T 2025', '+USD 2.060 millones', 'El superavit comercial de bienes no alcanza para equilibrar la cuenta corriente.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Servicios 1T 2025', '-USD 4.502 millones', 'Los servicios, en especial viajes personales, explican parte central de la salida de divisas.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Renta de inversion 1T 2025', '-USD 3.333 millones', 'Las remesas y rentas externas mantienen presion sobre el sector externo.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Exportaciones marzo 2026', 'USD 8.645 millones', 'Muestra capacidad exportadora, pero no elimina la restriccion externa.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Importaciones marzo 2026', 'USD 6.122 millones', 'La evolucion importadora incide sobre actividad, insumos y demanda de divisas.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Superavit comercial 1T 2026', 'USD 5.508 millones', 'Es una fortaleza coyuntural del frente comercial.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Inflacion 1T 2026', '9,4% acumulado; marzo 3,4% mensual', 'Afecta costos, salarios, precios, tasas y flujos nominales.', 'IE - Comercio Internacional Informe'),
    ('SSI exportaciones 2023', 'USD 2.449 millones', 'El SSI genera divisas de alto valor agregado.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('SSI en exportaciones de servicios', '15,2%', 'El sector tiene peso relevante dentro de servicios exportables.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Empleo SSI registrado 2023', '145.969 trabajadores', 'Evidencia base laboral calificada y formal.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Empresas SSI', '6.108 empresas', 'Muestra un ecosistema empresarial existente.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Remuneraciones SSI', '+78% vs. promedio privado formal', 'Refleja productividad y presion salarial relativa.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('Trabajadores IT con ingresos en USD', '32% total o parcialmente', 'Indica competencia global por talento y dolarizacion parcial.', 'Comercio Internacional / informe_salarios_argentina'),
    ('Brecha perfiles senior USD vs. pesos', 'hasta 62%', 'Aumenta costo de retencion de perfiles calificados.', 'IE - Comercio Internacional Informe / presentacion_slides'),
    ('I+D SSI', '3,6% de ventas; 14% de I+D empresaria nacional', 'El sector tiene intensidad innovadora superior al promedio.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ('ILIA 2024', 'Argentina 4° en LATAM; puntaje 47,4', 'Ubica capacidades en IA, pero tambien brechas de infraestructura y talento.', 'IE - Comercio Internacional Informe'),
    ('Economia regional comparada', 'PBI Argentina 2024 USD 638,37 mil millones; PBI per capita USD 13.969,8', 'Argentina conserva escala, pero queda por debajo de Chile y Uruguay en ingreso por habitante.', 'IE - TPI - PBI'),
]

factor_data = {
    'politicos': {
        'title': '2.1 Factores politicos',
        'situation': [
            'El entorno politico argentino se caracteriza por una tension entre politicas de promocion a la inversion y antecedentes de volatilidad institucional. Los materiales sobre instituciones remarcan que la confianza inversora puede verse afectada por rupturas de contratos, cambios regulatorios, alteraciones de reglas y baja previsibilidad en comparacion regional.',
            'Al mismo tiempo, los documentos de comercio internacional y SSI muestran instrumentos publicos que buscan mejorar las condiciones para invertir: Ley de Economia del Conocimiento, RIGI, RIMI y regimenes provinciales. Estos mecanismos no eliminan el riesgo argentino, pero funcionan como intentos de compensar incertidumbre mediante beneficios fiscales, estabilidad y tratamiento especifico para sectores estrategicos.',
        ],
        'data': [
            ('Ley de Economia del Conocimiento', 'Beneficios hasta el 31/12/2029; bono fiscal de hasta 70% de contribuciones patronales, ampliable al 80% en ciertos casos; reduccion de Ganancias 60%, 40% o 20% segun tamano de empresa.', 'Reduce costo laboral y fiscal en actividades intensivas en conocimiento.', 'IE - Comercio Internacional Informe'),
            ('RIGI', 'Estabilidad tributaria, aduanera y cambiaria por 30 anos para grandes proyectos; libre disponibilidad de divisas desde el tercer ano de adhesion.', 'Mejora previsibilidad para proyectos de gran escala.', 'IE - Comercio Internacional Informe / Comercio Internacional'),
            ('Promocion provincial', '23 de 24 provincias con legislacion propia; Cordoba con 10 anos de exenciones y promocion de empleo CTIM.', 'Permite elegir jurisdicciones con menor carga y mejor soporte sectorial.', 'IE - Comercio Internacional Informe'),
            ('Rupturas institucionales', 'Materiales citan corralito, estatizacion AFJP, intervencion INDEC y caso Tecpetrol como antecedentes.', 'Elevan percepcion de riesgo y tasa requerida por inversores.', 'PP IE / INFORME IE'),
        ],
        'impact': 'Para una empresa que evalua radicarse o expandirse, la politica incide sobre horizonte de inversion, costo de capital, repatriacion de utilidades, seleccion de provincia y confianza contractual. En un proyecto SSI, la estabilidad fiscal y regulatoria es especialmente importante porque los activos principales son talento, contratos, propiedad intelectual y continuidad operativa.',
        'opportunities': [
            'Usar regimenes de promocion para reducir costos fiscales y laborales.',
            'Elegir provincias con incentivos especificos para economia del conocimiento.',
            'Alinear el proyecto con objetivos publicos de generacion de divisas, empleo calificado e innovacion.',
        ],
        'risks': [
            'Cambios regulatorios que alteren supuestos de rentabilidad.',
            'Tramites y complejidad administrativa que demoren proyectos.',
            'Riesgo de que la estabilidad formal no compense la volatilidad macro-institucional.',
        ],
        'trends': [
            'Mayor competencia entre provincias para atraer empresas tecnologicas y empleo calificado.',
            'Necesidad de demostrar previsibilidad para que los regimenes promocionales sean creibles.',
            'Mayor peso de sectores generadores de divisas en la agenda de politica economica.',
        ],
        'partial': 'La dimension politica ofrece incentivos reales, pero Argentina sigue siendo un destino donde la decision de invertir exige cobertura institucional y contractual. La oportunidad no esta en ignorar el riesgo, sino en estructurar la inversion bajo marcos promocionales y con escenarios alternativos.',
    },
    'economicos': {
        'title': '2.2 Factores economicos',
        'situation': [
            'La economia argentina combina escala productiva, recursos y sectores dinamicos con restricciones macroeconomicas persistentes. Los materiales de PBI describen una economia con recuperaciones y caidas, baja inversion, inflacion, dificultades de credito y restriccion externa. Para un inversor, esto implica que la rentabilidad no depende solo de vender o producir, sino de poder financiarse, importar insumos, fijar precios y acceder a divisas.',
            'El eje mas fuerte de los documentos de comercio internacional es que el superavit comercial de bienes no alcanza por si solo. La cuenta corriente tambien depende de servicios, renta de inversion y financiamiento externo. Por eso, sectores como SSI son relevantes: exportan conocimiento, generan divisas y no dependen tanto de importaciones fisicas o del ciclo de commodities.',
        ],
        'data': [
            ('Cuenta corriente 1T 2025', '-USD 5.191 millones', 'Indica necesidad de financiamiento externo y fragilidad externa.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Bienes 1T 2025', '+USD 2.060 millones', 'El saldo positivo de bienes fue insuficiente frente a servicios y rentas.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Servicios 1T 2025', '-USD 4.502 millones', 'La salida por servicios presiona el equilibrio externo.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Superavit comercial 1T 2026', 'USD 5.508 millones', 'Fortaleza comercial reciente, pero no garantia estructural.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Exportaciones marzo 2026', 'USD 8.645 millones; +30,1% interanual', 'Muestra capacidad exportadora y mejora coyuntural.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Inflacion 1T 2026', '9,4% acumulado; marzo 3,4% mensual', 'Afecta costos, salarios, tasa de descuento y flujo de fondos.', 'IE - Comercio Internacional Informe'),
            ('PBI Argentina 2024', 'USD 638,37 mil millones; PBI per capita USD 13.969,8', 'Argentina mantiene escala regional, pero enfrenta problemas de productividad y estabilidad.', 'IE - TPI - PBI'),
        ],
        'impact': 'La dimension economica condiciona ingresos esperados, costos, capital de trabajo, financiamiento, tasa de descuento, VAN, TIR y Payback. En Argentina, un proyecto debe modelarse con escenarios de inflacion, tipo de cambio, disponibilidad de divisas, demanda interna y costos salariales. Para SSI, la oportunidad aumenta si los ingresos estan dolarizados por exportacion, pero los costos de talento tambien pueden dolarizarse parcial o informalmente.',
        'opportunities': [
            'Invertir en actividades que generen divisas o sustituyan importaciones criticas.',
            'Usar Argentina como plataforma de servicios exportables, especialmente en economia del conocimiento.',
            'Aprovechar costos relativos y talento local cuando el tipo de cambio y la estructura contractual lo permitan.',
        ],
        'risks': [
            'Inflacion que distorsione precios, salarios y proyecciones de flujo.',
            'Apreciacion cambiaria que reduzca competitividad exportadora.',
            'Restriccion externa que afecte importaciones, pagos al exterior o repatriacion de utilidades.',
            'Baja profundidad financiera y mayor tasa de descuento exigida por riesgo.',
        ],
        'trends': [
            'Mayor importancia de sectores exportadores de servicios para aliviar la restriccion externa.',
            'Necesidad de evaluar proyectos en moneda real y nominal para no confundir rentabilidad con inflacion.',
            'Mayor selectividad: no toda inversion es atractiva, pero las que generan divisas y productividad tienen mejor posicion.',
        ],
        'partial': 'La dimension economica es el centro del diagnostico. Argentina ofrece mercado, recursos y talento, pero el inversor debe evaluar la fragilidad externa y la volatilidad de precios. El SSI aparece como caso favorable porque convierte capital humano en exportaciones, aunque no queda aislado de los riesgos macro.',
    },
    'sociales': {
        'title': '2.3 Factores sociales',
        'situation': [
            'Los materiales muestran una estructura social dual. Argentina cuenta con capital humano calificado y polos de actividad tecnologica, pero tambien con pobreza, informalidad, desigualdad educativa y deterioros del salario real. Esta dualidad afecta tanto al consumo interno como a la disponibilidad futura de talento.',
            'En SSI, la base laboral es una fortaleza concreta: empleo registrado, empresas activas, salarios superiores al promedio privado y polos regionales. Sin embargo, la competencia global por talento, el trabajo remoto externo y la dolarizacion parcial de ingresos generan presion sobre costos y retencion.',
        ],
        'data': [
            ('Empleo SSI 2023', '145.969 trabajadores registrados', 'Base laboral formal y calificada.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Empresas SSI', '6.108 empresas', 'Ecosistema sectorial existente.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Concentracion geografica', 'CABA concentra 67,8% del empleo SSI; crecen Cordoba, Santa Fe, Mendoza y Rio Negro.', 'Permite pensar localizacion fuera de CABA.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Remuneraciones SSI', '+78% sobre promedio privado formal', 'Indica productividad y mayor costo relativo del talento.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Ingresos IT en USD', '32% cobra total o parcialmente en USD; senior en USD hasta 62% mas que en pesos.', 'Presiona retencion y politica salarial.', 'informe_salarios_argentina / IE - Comercio Internacional Informe'),
            ('Pobreza e informalidad', 'Los documentos de pobreza y PBI destacan pobreza, informalidad y brechas educativas como restricciones estructurales.', 'Limita mercado interno y pipeline de talento.', 'PDF - IE - TP - Pobreza / IE - TPI - PBI'),
        ],
        'impact': 'La dimension social afecta productividad, consumo, rotacion, capacitacion y escalabilidad. Para un proyecto orientado al mercado interno, pobreza e informalidad restringen demanda. Para un proyecto SSI exportador, el riesgo principal es la disponibilidad y retencion de perfiles calificados en un mercado laboral conectado con el exterior.',
        'opportunities': [
            'Aprovechar capital humano calificado en tecnologia, servicios profesionales, administracion e ingenieria.',
            'Desarrollar empleo formal en polos regionales con menor saturacion que CABA.',
            'Implementar capacitacion interna para ampliar la base de talento disponible.',
        ],
        'risks': [
            'Fuga de talento hacia empleo remoto internacional.',
            'Aumento de costos salariales en perfiles senior o especializados.',
            'Brechas educativas y sociales que limiten la expansion futura del sector.',
        ],
        'trends': [
            'Mayor competencia por perfiles bilingues, senior y vinculados a IA.',
            'Crecimiento de modalidades de compensacion flexibles o parcialmente dolarizadas.',
            'Necesidad de politicas de formacion para que el crecimiento tecnologico sea escalable.',
        ],
        'partial': 'Socialmente, Argentina tiene una ventaja fuerte en talento, pero no es infinita ni automatica. Para invertir en SSI, la estrategia laboral debe ser tan importante como la estrategia comercial: retener talento, formar perfiles y elegir ubicacion son decisiones economicas centrales.',
    },
    'tecnologicos': {
        'title': '2.4 Factores tecnologicos',
        'situation': [
            'La dimension tecnologica es una de las principales oportunidades de Argentina como destino de inversion. Los materiales muestran capacidades en software, servicios informaticos, IA, fintech, salud, AgTech, e-commerce, IoT y actividades basadas en conocimiento. Estas areas pueden elevar productividad y diversificar exportaciones.',
            'El limite aparece en infraestructura digital, conectividad, 5G, dispositivos, ciberseguridad, patentes, inversiones entrantes y alfabetizacion en IA. Por lo tanto, Argentina no debe presentarse como un ecosistema tecnologico sin problemas, sino como un pais con capacidades reales y brechas de adopcion.',
        ],
        'data': [
            ('Exportaciones SSI 2023', 'USD 2.449 millones; 15,2% de exportaciones de servicios', 'Sector tecnologico con capacidad de generar divisas.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Crecimiento SSI 2015-2022', '+10,2% anual', 'Muestra dinamismo previo y potencial exportador.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Participacion mundial Argentina', '0,59% en 2011 a 0,27%/0,3% en 2023', 'Indica perdida relativa de posicion global.', 'IE - Comercio Internacional Informe'),
            ('I+D SSI', '3,6% de ventas; 14% de I+D empresaria nacional', 'Mayor intensidad innovadora que otros sectores.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('ILIA 2024', 'Argentina 47,4; 4° LATAM; promedio regional 40,3', 'Capacidad regional relevante en IA, pero por debajo de Chile, Uruguay y Brasil.', 'IE - Comercio Internacional Informe'),
            ('IA y empleo formal', '54% del empleo formal privado en ocupaciones con al menos la mitad de tareas ejecutables por IA generativa; 71% con relacion de complementariedad.', 'La IA puede aumentar productividad, pero exige adaptacion laboral.', 'IE - Comercio Internacional Informe'),
        ],
        'impact': 'La tecnologia impacta sobre productividad, diferenciacion, escalabilidad y costos. Para una empresa internacional, Argentina puede servir como base de desarrollo, soporte, automatizacion, analitica o productos digitales. Pero proyectos intensivos en conectividad, datos o IA deben evaluar infraestructura, disponibilidad de equipos, ciberseguridad y capacidad de retencion de talento.',
        'opportunities': [
            'Invertir en verticales de alto valor: IA, fintech, healthtech, AgTech, e-commerce e IoT industrial.',
            'Aprovechar instrumentos de I+D+i como FONSOFT, FONTAR y FONARSEC mencionados en los materiales.',
            'Conectar software con sectores tradicionales para mejorar productividad exportadora.',
        ],
        'risks': [
            'Rezago 5G y conectividad insuficiente para IA en tiempo real, IoT o robotica industrial.',
            'Costo de dispositivos y equipamiento que limite adopcion tecnologica.',
            'Brechas STEM, ciberseguridad y baja cantidad de patentes o inversiones entrantes.',
        ],
        'trends': [
            'Mayor demanda de soluciones de IA generativa y automatizacion de procesos.',
            'Creciente necesidad de ciberseguridad y gobernanza tecnologica.',
            'Digitalizacion de actividades tradicionales como agro, salud, comercio, energia y administracion.',
        ],
        'partial': 'Tecnologicamente, Argentina tiene una oportunidad clara: usar conocimiento para aumentar productividad y exportaciones. El riesgo es que el talento exista, pero la infraestructura y el financiamiento no permitan escalar al ritmo de la demanda global.',
    },
    'legales': {
        'title': '2.5 Factores legales',
        'situation': [
            'El entorno legal argentino combina incentivos especificos con alta complejidad. Los materiales identifican seguridad juridica, estabilidad regulatoria, presion tributaria, normas laborales, restricciones cambiarias y eficacia judicial como factores que condicionan inversion y formalizacion.',
            'Para SSI y economia del conocimiento, el marco legal ofrece beneficios concretos. Sin embargo, el cumplimiento tributario, el Convenio Multilateral, la litigiosidad, las normas laborales y la historia de cambios regulatorios obligan a estructurar cuidadosamente cualquier proyecto.',
        ],
        'data': [
            ('Ley de Economia del Conocimiento', 'Bono fiscal de contribuciones; reduccion de Ganancias; vigencia hasta 2029.', 'Reduce costos si la empresa cumple requisitos de acceso.', 'IE - Comercio Internacional Informe'),
            ('Exportaciones de servicios SSI', 'Exentas de IIBB en jurisdicciones clave segun los materiales.', 'Preserva competitividad de ventas externas.', 'IE - Comercio Internacional Informe'),
            ('Cordoba Ley 10.722', 'Exencion 100% en IIBB por 10 anos y promocion de empleo CTIM.', 'Incentivo local para radicacion tecnologica.', 'IE - Comercio Internacional Informe'),
            ('RIGI', 'Estabilidad 30 anos; beneficios tributarios, aduaneros y cambiarios.', 'Puede mejorar viabilidad de grandes proyectos.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
            ('Marco tributario 2026', 'Ganancias sociedades 25-31,5%; IVA 21%; IIBB 3-7%; municipal 0,5-1%; debitos y creditos 0,6%.', 'Afecta estructura de costos y formalizacion.', 'PP IE'),
            ('RIMI', 'PyMEs USD 150K-9M; amortizacion acelerada y devolucion de IVA en 3 meses.', 'Herramienta posible para inversiones menores que RIGI.', 'PP IE'),
        ],
        'impact': 'La dimension legal define como se opera: estructura societaria, impuestos, contratacion, exportaciones, importaciones, facturacion, propiedad intelectual, giro de utilidades y defensa contractual. En SSI, tambien afecta proteccion de software, contratacion remota, tributacion interprovincial y cobro de exportaciones.',
        'opportunities': [
            'Aprovechar beneficios legales y fiscales de economia del conocimiento.',
            'Localizar operaciones en provincias con mejor combinacion de incentivos, talento y costos.',
            'Usar regimenes de inversion para mejorar flujo de fondos, amortizacion y disponibilidad de divisas.',
        ],
        'risks': [
            'Complejidad tributaria nacional, provincial y municipal.',
            'Riesgo laboral y litigiosidad que incrementen costos no previstos.',
            'Normativa cambiaria o contractual cambiante que afecte pagos, cobros o repatriacion.',
        ],
        'trends': [
            'Mayor utilizacion de regimenes promocionales para atraer inversion.',
            'Necesidad de ordenar tributacion de trabajo remoto y operaciones multiprovinciales.',
            'Mayor demanda de seguridad juridica para proyectos de largo plazo.',
        ],
        'partial': 'Legalmente, Argentina ofrece instrumentos valiosos, pero exige gestion profesional de cumplimiento. La inversion viable no es la que ignora la complejidad, sino la que la incorpora en sus flujos, contratos y decisiones de localizacion.',
    },
}


def add_factor_section(doc, factor):
    add_h2(doc, factor['title'])
    add_h3(doc, 'Situacion actual')
    for paragraph in factor['situation']:
        add_p(doc, paragraph)
    add_h3(doc, 'Datos cuantitativos y evidencia documental')
    add_table(doc, ['Tema', 'Dato en los materiales', 'Lectura para inversion', 'Fuente entregada'], factor['data'], [1.55, 1.7, 2.05, 1.0])
    add_h3(doc, 'Impacto sobre empresas e inversiones')
    add_p(doc, factor['impact'])
    add_h3(doc, 'Oportunidades')
    for item in factor['opportunities']:
        add_bullet(doc, item)
    add_h3(doc, 'Riesgos')
    for item in factor['risks']:
        add_bullet(doc, item)
    add_h3(doc, 'Tendencias futuras')
    for item in factor['trends']:
        add_bullet(doc, item)
    add_h3(doc, 'Conclusion parcial')
    add_callout(doc, 'Lectura del factor', factor['partial'], fill=PALE_GREEN)


def build():
    doc = setup_doc()
    add_title(doc)
    page_break(doc)

    add_h1(doc, 'RESUMEN')
    add_p(doc, 'El presente informe analiza el estado actual de Argentina como destino de inversion empresarial mediante un enfoque PESTEL adaptado a los documentos de la materia Ingenieria Economica. El diagnostico muestra un pais con capacidades productivas reales, sectores exportadores, talento calificado, economia del conocimiento y regimenes de promocion; pero tambien con restricciones macroeconomicas, volatilidad institucional, pobreza, informalidad, baja profundidad financiera, presion tributaria y complejidad regulatoria.')
    add_p(doc, 'El informe no interpreta a Argentina como un destino de inversion homogeneo. La conclusion central es que el pais puede ser atractivo de manera selectiva: especialmente para proyectos que generen divisas, eleven productividad, aprovechen talento local, incorporen tecnologia y operen bajo marcos legales promocionales. En cambio, las inversiones muy dependientes del mercado interno, de importaciones criticas o de reglas cambiarias estables enfrentan mayor riesgo.')
    add_p(doc, 'El sector de Software y Servicios Informaticos se utiliza como caso testigo porque los materiales lo presentan como una actividad capaz de transformar capital humano en exportaciones de alto valor agregado. En 2023 el SSI exporto USD 2.449 millones, represento 15,2% de las exportaciones de servicios y empleo 145.969 trabajadores registrados. Estos datos permiten conectar comercio internacional, restriccion externa, productividad, salarios, innovacion y evaluacion de proyectos.')
    add_p(doc, 'La lectura tecnica se apoya en indicadores cuantitativos extraidos de las presentaciones e informes entregados: cuenta corriente, saldo de bienes, deficit de servicios, exportaciones, inflacion, PBI, empleo SSI, empresas, salarios, I+D, IA, regimenes legales y herramientas de evaluacion financiera como VAN, TIR, ROI y Payback. No se incorporan datos externos nuevos.')
    add_table(doc, ['Eje del diagnostico', 'Hallazgo principal', 'Implicancia para inversion'], [
        ('Sector externo', 'El superavit de bienes no alcanza para equilibrar cuenta corriente por servicios y renta de inversion.', 'Conviene priorizar actividades que generen divisas y reduzcan dependencia externa.'),
        ('Macroeconomia', 'Inflacion, tipo de cambio, credito y restriccion externa elevan incertidumbre.', 'Los proyectos requieren escenarios, tasa de descuento prudente y gestion de capital de trabajo.'),
        ('SSI', 'Sector con exportaciones, empleo formal, salarios altos e I+D.', 'Puede funcionar como plataforma exportadora si retiene talento y mantiene competitividad.'),
        ('Marco legal', 'Existen incentivos, pero tambien complejidad tributaria y regulatoria.', 'La localizacion y el cumplimiento son parte de la estrategia economica.'),
    ], [1.6, 2.9, 1.8], caption='Sintesis ejecutiva del informe')
    page_break(doc)

    add_h1(doc, '1. INTRODUCCION')
    add_p(doc, 'La decision de invertir en un pais requiere analizar mas que una oportunidad comercial puntual. Desde Ingenieria Economica, una inversion implica sacrificar recursos presentes con la expectativa de obtener beneficios futuros bajo riesgo e incertidumbre. Por eso, antes de calcular indicadores como VAN, TIR, ROI o Payback, es necesario comprender el entorno que condiciona los flujos de fondos.')
    add_p(doc, 'Este informe aplica un esquema PESTEL al caso argentino, tomando como base exclusiva los materiales entregados por la catedra y por los grupos de trabajo. El objetivo es ordenar factores politicos, economicos, sociales, tecnologicos y legales que afectan la inversion empresarial. La dimension ambiental no se desarrolla por indicacion expresa de trabajo.')
    add_p(doc, 'El enfoque elegido combina dos niveles. Primero, se analiza Argentina como destino de inversion en sentido general: estabilidad, PBI, comercio exterior, inflacion, pobreza, informalidad, tecnologia y seguridad juridica. Segundo, se utiliza el sector SSI como caso testigo, porque permite observar una oportunidad concreta en economia del conocimiento: exportar servicios, generar divisas, emplear talento calificado y operar con menor dependencia de insumos fisicos.')
    add_p(doc, 'El informe busca responder una pregunta practica: bajo que condiciones Argentina puede ser atractiva para invertir. La respuesta no es absoluta. Argentina no aparece como un destino de bajo riesgo, pero si como un pais con oportunidades selectivas cuando el proyecto esta bien estructurado, tiene cobertura frente a variables macroeconomicas y se apoya en sectores con capacidad exportadora o productividad creciente.')
    add_h2(doc, 'Metodologia y alcance')
    for item in [
        'Se utilizaron unicamente documentos entregados: presentaciones, informes PDF, archivos de salarios, comercio internacional, PBI, pobreza, inversiones, instituciones y SSI.',
        'Los datos cuantitativos se incorporaron solo cuando aparecian en esos materiales o en sus fuentes citadas internamente.',
        'Se evitaron datos externos nuevos, busquedas web y afirmaciones no trazables a los archivos de trabajo.',
        'Se adapto el esquema PESTEL excluyendo el apartado ambiental, para mantener el criterio indicado.',
    ]:
        add_bullet(doc, item)
    add_h2(doc, 'Estructura del informe')
    for item in [
        'Resumen tecnico con hallazgos principales.',
        'Introduccion con alcance, metodologia y enfoque de inversion.',
        'Analisis y desarrollo con factores politicos, economicos, sociales, tecnologicos y legales.',
        'Matriz integrada, escenarios estrategicos 2026 y relacion explicita con contenidos de Ingenieria Economica.',
        'Conclusiones, bibliografia, fuentes de datos y glosario.',
    ]:
        add_number(doc, item)
    page_break(doc)

    add_h1(doc, '2. ANALISIS Y DESARROLLO')
    add_h2(doc, '2.0 Indicadores de referencia')
    add_p(doc, 'Antes de desarrollar cada factor, conviene ordenar los principales datos cuantitativos que sostienen el diagnostico. Estos indicadores permiten pasar de una lectura puramente cualitativa a una evaluacion tecnica del entorno de inversion.')
    add_table(doc, ['Indicador', 'Valor en los materiales', 'Interpretacion para inversion', 'Fuente entregada'], quant_rows, [1.35, 1.45, 2.55, 1.15], caption='Indicadores cuantitativos utilizados en el informe')
    add_p(doc, 'La lectura conjunta de estos datos muestra una economia con capacidad exportadora, escala regional y sectores de alto valor, pero con vulnerabilidades persistentes. La restriccion externa aparece como un condicionante central: incluso con superavit comercial de bienes, los egresos por servicios y renta pueden mantener deficit de cuenta corriente.')

    for key in ['politicos', 'economicos', 'sociales', 'tecnologicos', 'legales']:
        add_factor_section(doc, factor_data[key])
        page_break(doc)

    add_h2(doc, '2.6 Matriz integrada de impacto para la inversion')
    add_table(doc, ['Factor', 'Oportunidades', 'Riesgos', 'Impacto sobre inversion'], [
        ('Politico', 'Regimenes de promocion, estabilidad fiscal sectorial, incentivos nacionales y provinciales.', 'Volatilidad institucional, cambios regulatorios, tramites y antecedentes contractuales.', 'Alto: define horizonte, confianza, tasa de descuento y decision de radicacion.'),
        ('Economico', 'Sectores exportadores, recursos productivos, servicios basados en conocimiento y capacidad de generar divisas.', 'Inflacion, restriccion externa, apreciacion cambiaria, bajo credito y costos inciertos.', 'Alto: afecta ingresos, costos, financiamiento, VAN, TIR y Payback.'),
        ('Social', 'Capital humano calificado, empleo SSI formal, polos regionales y salarios asociados a productividad.', 'Pobreza, informalidad, brechas educativas, fuga de talento y presion salarial.', 'Medio/alto: condiciona consumo, productividad, contratacion y escalabilidad.'),
        ('Tecnologico', 'IA, I+D, software, fintech, salud, AgTech, e-commerce e IoT.', 'Infraestructura digital insuficiente, rezago 5G, ciberseguridad, patentes y brechas STEM.', 'Alto: define productividad, diferenciacion y competitividad internacional.'),
        ('Legal', 'Ley de Economia del Conocimiento, RIGI, RIMI, exenciones provinciales y beneficios tributarios.', 'Presion tributaria, costos laborales, restricciones cambiarias, litigiosidad y Convenio Multilateral.', 'Alto: determina costos de cumplimiento, formalizacion y previsibilidad contractual.'),
    ], [1.05, 2.0, 2.0, 1.25], caption='Matriz resumen del analisis PESTEL adaptado')
    add_p(doc, 'La matriz permite observar que las oportunidades mas relevantes se concentran en sectores exportadores, tecnologicos y promocionados. Los riesgos, en cambio, son transversales: macroeconomia, instituciones, impuestos, trabajo y divisas. Por eso, la decision de inversion no debe basarse solo en el atractivo sectorial, sino en la capacidad del proyecto para soportar escenarios adversos.')

    add_h2(doc, '2.7 Diagnostico estrategico')
    add_callout(doc, 'Diagnostico central', 'Argentina es un destino de inversion con capacidades reales, pero de riesgo elevado. La oportunidad aparece cuando el proyecto combina generacion de divisas, talento local, tecnologia, productividad e incentivos legales. El riesgo aumenta cuando la inversion depende principalmente del mercado interno, del credito local barato, de importaciones restringidas o de estabilidad regulatoria no asegurada.', fill=PALE_GOLD)
    add_h3(doc, 'Fortalezas estrategicas')
    for item in [
        'Escala economica regional y estructura productiva diversificada.',
        'Capacidad exportadora en bienes y servicios.',
        'Talento calificado en SSI y economia del conocimiento.',
        'Regimenes de promocion nacionales y provinciales.',
        'Potencial de aumentar productividad mediante tecnologia e IA.',
    ]:
        add_bullet(doc, item)
    add_h3(doc, 'Debilidades estructurales')
    for item in [
        'Restriccion externa y fragilidad de cuenta corriente.',
        'Inflacion y volatilidad de precios relativos.',
        'Baja profundidad financiera y costo de capital elevado.',
        'Pobreza, informalidad y brechas educativas.',
        'Complejidad tributaria, laboral y administrativa.',
    ]:
        add_bullet(doc, item)
    add_h3(doc, 'Lectura para SSI')
    add_p(doc, 'El SSI es uno de los sectores que mejor encaja con las necesidades macroeconomicas argentinas porque genera divisas, utiliza talento, demanda menos insumos fisicos importados y puede vender al exterior. Sin embargo, tambien enfrenta riesgos propios: perdida de participacion global, atraso de infraestructura, presion salarial, fuga de talento y dependencia de reglas cambiarias previsibles.')
    page_break(doc)

    add_h2(doc, '2.8 Escenarios estrategicos para 2026')
    add_p(doc, 'Tomando la estructura del informe del Octa como mejora formal, se incorporan escenarios estrategicos. Estos escenarios no agregan datos externos: ordenan posibles resultados a partir de las variables presentes en los materiales entregados.')
    add_table(doc, ['Escenario', 'Supuestos principales', 'Resultado esperado para inversion', 'Estrategia recomendada'], [
        ('Base', 'Se mantienen superavit comercial, inflacion relevante, restriccion externa y uso de incentivos sectoriales.', 'Argentina es atractiva de forma selectiva, especialmente en proyectos exportadores y tecnologicos.', 'Invertir por etapas, proteger flujo de fondos, usar regimenes promocionales y priorizar ingresos en divisas.'),
        ('Favorable', 'Mejora la estabilidad macro, se reduce incertidumbre regulatoria y crecen exportaciones de servicios.', 'Aumenta la viabilidad de proyectos de mayor escala y plazos mas largos.', 'Expandir operaciones, contratar talento, invertir en I+D y consolidar polos regionales.'),
        ('Riesgo', 'Apreciacion cambiaria, inflacion persistente, restricciones externas o cambios normativos afectan competitividad.', 'Sube la tasa de descuento, se alarga Payback y pueden deteriorarse VAN/TIR.', 'Reducir CAPEX inicial, usar contratos flexibles, diversificar mercados y mantener cobertura cambiaria.'),
    ], [1.0, 2.05, 1.75, 1.45], caption='Escenarios de decision para inversion')

    add_h2(doc, '2.9 Recomendaciones estrategicas por tipo de inversion')
    add_table(doc, ['Tipo de inversion', 'Atractivo en Argentina', 'Condiciones de viabilidad', 'Riesgo principal'], [
        ('SSI exportador', 'Alto', 'Ingresos externos, talento retenible, uso de Ley EdC, infraestructura digital y contratos internacionales.', 'Fuga de talento, costos dolarizados y apreciacion cambiaria.'),
        ('Servicios profesionales exportables', 'Medio/alto', 'Capacidad bilingue, procesos digitales, demanda externa y bajo CAPEX fisico.', 'Competencia regional y volatilidad macro.'),
        ('Industria mercado interno', 'Medio', 'Demanda estable, financiamiento, insumos disponibles y precios actualizables.', 'Inflacion, caida de consumo y restricciones de importacion.'),
        ('Grandes proyectos de infraestructura/energia', 'Medio/alto', 'RIGI, escala, financiamiento externo y estabilidad contractual.', 'Riesgo institucional, plazos largos y cambios regulatorios.'),
        ('Startups tecnologicas', 'Medio', 'Talento, nicho global, acceso a capital y proteccion de propiedad intelectual.', 'Falta de financiamiento local y migracion de emprendedores.'),
    ], [1.3, 1.0, 2.55, 1.45], caption='Recomendaciones segun perfil de proyecto')

    add_h2(doc, '2.10 Relacion explicita con contenidos de Ingenieria Economica')
    add_p(doc, 'El analisis PESTEL no reemplaza la evaluacion financiera, sino que la alimenta. Cada factor del entorno modifica supuestos del flujo de fondos: inversion inicial, ingresos, costos, capital de trabajo, tasa de descuento, riesgo, valor residual y plazo de recupero.')
    add_table(doc, ['Contenido de IE', 'Aplicacion al caso Argentina', 'Efecto sobre la decision'], [
        ('Flujo de fondos', 'Debe incorporar inflacion, tipo de cambio, costos laborales, impuestos, beneficios fiscales e ingresos externos.', 'Permite estimar si el proyecto genera fondos suficientes bajo escenarios realistas.'),
        ('Valor del dinero en el tiempo', 'La inflacion y el costo de oportunidad hacen necesario comparar flujos en terminos homogeneos.', 'Evita confundir crecimiento nominal con creacion real de valor.'),
        ('Tasa de descuento', 'Debe reflejar riesgo macro, legal, cambiario, sectorial y financiero.', 'A mayor riesgo, mayor tasa requerida y menor VAN.'),
        ('VAN', 'Mide si el proyecto genera valor presente despues de descontar la inversion inicial.', 'Si el VAN es positivo bajo escenarios prudentes, la inversion es mas defendible.'),
        ('TIR', 'Permite comparar rentabilidad del proyecto contra tasa requerida.', 'Si la TIR supera el costo de capital ajustado por riesgo, el proyecto puede aceptarse.'),
        ('ROI', 'Relaciona ganancia con monto invertido, aunque no considera tiempo.', 'Sirve como indicador rapido, pero debe complementarse con VAN y TIR.'),
        ('Payback', 'Mide recupero de inversion, clave en economias volatiles.', 'Cuanto mas corto el Payback, menor exposicion a cambios macro y regulatorios.'),
        ('Financiamiento', 'Sistemas frances, aleman, americano y leasing alteran cuotas, intereses y liquidez.', 'La estructura de deuda puede mejorar o deteriorar flujo y riesgo.'),
    ], [1.3, 3.0, 2.0], caption='Vinculo entre PESTEL y evaluacion de proyectos')
    add_p(doc, 'En sintesis, los factores PESTEL se traducen en variables financieras. La inflacion modifica flujos nominales; el tipo de cambio afecta ingresos exportadores y costos importados; la seguridad juridica incide en tasa de descuento; la disponibilidad de talento afecta salarios y productividad; y los beneficios legales modifican impuestos, CAPEX, capital de trabajo o recupero de IVA.')

    add_h2(doc, '2.11 Caso testigo: SSI como plataforma de inversion')
    add_p(doc, 'El SSI funciona en este informe como caso testigo porque conecta casi todos los problemas y oportunidades del entorno argentino. Desde el punto de vista macroeconomico, aporta divisas; desde el punto de vista social, emplea talento calificado; desde el punto de vista tecnologico, invierte en I+D; y desde el punto de vista legal, puede acceder a regimenes de promocion especificos.')
    add_p(doc, 'La ventaja principal del SSI frente a sectores mas intensivos en bienes fisicos es que puede exportar conocimiento sin depender de grandes importaciones de insumos. Esto es especialmente importante en un pais con restriccion externa, porque una actividad que genera dolares y demanda relativamente menos divisas para producir puede mejorar la sostenibilidad del crecimiento.')
    add_table(doc, ['Variable SSI', 'Dato usado', 'Lectura economica', 'Fuente entregada'], [
        ('Exportaciones', 'USD 2.449 millones en 2023', 'Capacidad de generar divisas de alto valor agregado.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
        ('Peso en servicios', '15,2% de exportaciones de servicios', 'Relevancia dentro de una cuenta de servicios que suele ser deficitaria.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
        ('Empleo', '145.969 registrados', 'Base formal de capital humano calificado.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
        ('Empresas', '6.108 firmas', 'Ecosistema empresarial existente para radicacion, proveedores y alianzas.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
        ('Salarios', '+78% vs. promedio privado formal', 'Senal de productividad, pero tambien de costo laboral especializado.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
        ('I+D', '3,6% de ventas y 14% de I+D empresaria nacional', 'Mayor intensidad innovadora y potencial de diferenciacion.', 'Comercio Internacional / IE - Comercio Internacional Informe'),
    ], [1.25, 1.35, 2.55, 1.15], caption='SSI como caso testigo de inversion')
    add_h3(doc, 'Condiciones para que el SSI sea una oportunidad real')
    for item in [
        'Sostener competitividad externa mediante contratos en divisas, calidad tecnica y gestion del tipo de cambio.',
        'Retener talento con esquemas salariales competitivos, capacitacion y carrera profesional.',
        'Usar beneficios de Economia del Conocimiento y promociones provinciales para reducir costos.',
        'Invertir en infraestructura digital, ciberseguridad, procesos de calidad e I+D.',
        'Evitar depender exclusivamente del mercado interno cuando el objetivo es diversificar divisas.',
    ]:
        add_bullet(doc, item)
    page_break(doc)

    add_h2(doc, '2.12 Criterios de decision para un inversor')
    add_p(doc, 'A partir del diagnostico PESTEL, un inversor no deberia preguntar solamente si Argentina es barata o si tiene talento. La pregunta correcta es si el proyecto puede convertir esas ventajas en flujos de fondos estables bajo escenarios de volatilidad.')
    add_table(doc, ['Criterio', 'Pregunta tecnica', 'Senal favorable', 'Senal de alerta'], [
        ('Divisas', 'El proyecto genera ingresos externos?', 'Exporta servicios o bienes con baja demanda de importaciones.', 'Depende de pesos y necesita importar insumos criticos.'),
        ('Talento', 'Puede contratar y retener perfiles clave?', 'Tiene plan salarial, capacitacion y ubicacion con oferta laboral.', 'Compite por seniority escaso sin estrategia de retencion.'),
        ('Legal', 'Puede acceder a incentivos y cumplir requisitos?', 'Usa EdC, RIGI, RIMI o promocion provincial con asesoramiento.', 'Subestima impuestos, Convenio Multilateral o restricciones cambiarias.'),
        ('Finanzas', 'El VAN resiste una tasa de descuento alta?', 'Payback razonable y sensibilidad positiva.', 'Rentabilidad solo aparece en escenario optimista.'),
        ('Tecnologia', 'La infraestructura soporta la operacion?', 'Conectividad, ciberseguridad y herramientas adecuadas.', 'Requiere 5G, equipos o datos que no estan garantizados.'),
        ('Mercado', 'La demanda es suficientemente diversificada?', 'Clientes externos o cartera no concentrada.', 'Dependencia de mercado interno debil o pocos clientes.'),
    ], [1.05, 1.75, 1.85, 1.65], caption='Checklist tecnico para decidir inversion')
    add_p(doc, 'Este criterio permite pasar del diagnostico general a una decision de proyecto. Un mismo pais puede ser riesgoso para una inversion y atractivo para otra. La clave es medir exposicion: cuanto mas dependa el proyecto de variables inestables, mayor debera ser la tasa requerida y mas corto deberia ser el plazo de recupero buscado.')
    page_break(doc)

    add_h2(doc, '2.13 Lectura ampliada de variables cuantitativas')
    add_h3(doc, 'Sector externo y restriccion de divisas')
    add_p(doc, 'Los datos de balanza de pagos muestran que la restriccion externa no se explica solamente por importaciones de bienes. En 1T 2025, el saldo de bienes fue positivo, pero la cuenta corriente cerro en deficit por servicios y renta. Para una empresa, esto significa que la disponibilidad de divisas puede depender de factores macro que exceden su propio desempeno operativo.')
    add_p(doc, 'Esta lectura es importante para proyectos que pagan licencias, royalties, equipamiento, deuda o dividendos al exterior. Aun cuando una empresa sea rentable en pesos, puede enfrentar restricciones si necesita acceder a moneda extranjera. Por eso, los proyectos exportadores tienen una posicion relativa mejor: generan su propia fuente de divisas.')
    add_h3(doc, 'Inflacion, tasa y flujo de fondos')
    add_p(doc, 'La inflacion acumulada mencionada en los materiales para 1T 2026 obliga a trabajar con flujos consistentes. Si se proyectan ingresos nominales, tambien deben proyectarse costos nominales, impuestos y capital de trabajo. Si se trabaja en moneda real, la tasa de descuento debe ser real. Mezclar flujos nominales con tasas reales puede sobreestimar o subestimar el VAN.')
    add_p(doc, 'En economias volatiles, el Payback gana importancia practica porque mide exposicion temporal. Un proyecto con VAN positivo pero recupero muy largo puede ser fragil frente a cambios regulatorios, cambiarios o salariales. Por eso, el analisis financiero debe incluir sensibilidad ante inflacion, tipo de cambio y costo de capital.')
    page_break(doc)

    add_h2(doc, '2.14 Lectura ampliada de talento, tecnologia y legalidad')
    add_h3(doc, 'Talento como activo economico')
    add_p(doc, 'El capital humano no aparece solo como variable social, sino como activo economico. En SSI, los salarios superiores al promedio privado formal pueden interpretarse como senal de productividad y de integracion con mercados internacionales. Pero tambien implican que el costo laboral puede moverse con logicas distintas al salario promedio de la economia.')
    add_p(doc, 'La dolarizacion parcial de ingresos IT y la brecha de perfiles senior muestran que la empresa debe presupuestar retencion. Si el flujo de fondos subestima salarios o rotacion, el proyecto puede parecer rentable en papel y perder margen en la ejecucion.')
    add_h3(doc, 'Tecnologia y productividad')
    add_p(doc, 'La tecnologia es el puente entre inversion y crecimiento potencial. Los materiales muestran oportunidades en IA, fintech, salud, AgTech, e-commerce e IoT. Sin embargo, la adopcion efectiva depende de infraestructura, conectividad, talento STEM, ciberseguridad y financiamiento. Por eso, la tecnologia no debe tratarse como una promesa abstracta, sino como una variable operacional que requiere CAPEX, OPEX y gestion.')
    add_h3(doc, 'Legalidad como costo y como oportunidad')
    add_p(doc, 'El marco legal argentino tiene doble efecto. Por un lado, aumenta costos de cumplimiento por impuestos, normas laborales, Convenio Multilateral y riesgo regulatorio. Por otro, puede mejorar el proyecto si se aprovechan beneficios como EdC, RIGI, RIMI o promociones provinciales. Para Ingenieria Economica, esto implica que los incentivos deben modelarse en el flujo de fondos, no mencionarse solo como contexto.')
    page_break(doc)

    add_h2(doc, '2.15 Sintesis final del analisis')
    add_p(doc, 'La lectura integrada permite ubicar a Argentina en una categoria intermedia: no es un destino de inversion simple, pero tampoco es un pais sin oportunidades. La inversion mas defendible es aquella que transforma ventajas locales en ingresos externos, productividad o ahorro de costos medible.')
    add_p(doc, 'El SSI resume esa logica. Tiene talento, empresas, exportaciones, salarios de productividad e I+D. Al mismo tiempo, enfrenta riesgos de competitividad, infraestructura, regulacion y talento. Por eso, el informe no recomienda invertir en Argentina de forma indiscriminada, sino disenar proyectos selectivos, con cobertura y foco en sectores donde el pais tiene capacidades demostradas.')
    add_table(doc, ['Decision', 'Conviene cuando', 'No conviene cuando'], [
        ('Invertir', 'El proyecto exporta, genera divisas, usa talento local, tiene incentivos y Payback razonable.', 'La rentabilidad depende solo de estabilidad macro o regulatoria no asegurada.'),
        ('Postergar', 'Faltan definiciones regulatorias, financiamiento o estructura salarial.', 'El costo de esperar es menor que el riesgo de entrar mal estructurado.'),
        ('Entrar por etapas', 'El mercado es atractivo pero las variables macro siguen inestables.', 'Se puede validar demanda y talento antes de comprometer CAPEX alto.'),
        ('Reubicar provincia', 'Hay diferencias relevantes de impuestos, talento o beneficios locales.', 'La ubicacion elegida encarece cumplimiento o dificulta contratacion.'),
    ], [1.2, 2.55, 2.55], caption='Sintesis de decision estrategica')
    page_break(doc)

    add_h1(doc, '3. CONCLUSIONES')
    add_p(doc, 'Argentina puede ser atractiva como destino de inversion, pero no de manera general ni automatica. El pais posee capacidades productivas, sectores exportadores, recursos, talento, base tecnologica y regimenes promocionales. Sin embargo, esas fortalezas conviven con problemas persistentes: inflacion, restriccion externa, bajo credito, informalidad, pobreza, volatilidad normativa, complejidad tributaria y riesgo institucional.')
    add_p(doc, 'La principal conclusion tecnica es que el inversor debe mirar Argentina con una estrategia selectiva. Los proyectos mas convenientes son aquellos que generan divisas, incorporan tecnologia, elevan productividad, aprovechan talento local, reducen dependencia de importaciones fisicas y pueden operar bajo marcos promocionales. En esta categoria, el SSI aparece como uno de los mejores ejemplos del material analizado.')
    add_p(doc, 'El sector SSI muestra datos concretos: exportaciones por USD 2.449 millones en 2023, 15,2% de las exportaciones de servicios, 145.969 empleos registrados, 6.108 empresas, remuneraciones 78% superiores al promedio privado formal e inversion en I+D equivalente al 3,6% de sus ventas. Estos indicadores justifican utilizarlo como caso testigo de economia del conocimiento, no como unica actividad posible.')
    add_p(doc, 'Aun asi, el SSI no esta libre de riesgos. La perdida de participacion mundial, el rezago 5G, las brechas de infraestructura digital, la competencia por talento, la dolarizacion parcial de salarios y la incertidumbre cambiaria pueden afectar margenes y escalabilidad. Por eso, una empresa de software que invierta en Argentina deberia priorizar exportaciones, retencion de talento, estructura legal eficiente, localizacion provincial optimizada y escenarios financieros conservadores.')
    add_callout(doc, 'Respuesta final', 'Argentina es un destino de inversion posible y con oportunidades relevantes, pero exige gestion de riesgo. No conviene presentarla como un pais de bajo riesgo; conviene presentarla como un pais con capacidades reales donde la inversion es atractiva si se orienta a sectores productivos, tecnologicos y exportadores, con cobertura macroeconomica y legal.', fill=PALE_GREEN)
    add_h2(doc, 'Condiciones recomendadas para invertir')
    for item in [
        'Priorizar ingresos en divisas o mercados externos.',
        'Usar regimenes legales y fiscales disponibles: Ley de Economia del Conocimiento, RIGI, RIMI y beneficios provinciales.',
        'Evaluar proyectos con escenarios de inflacion, tipo de cambio y tasa de descuento.',
        'Reducir exposicion inicial mediante inversion por etapas y Payback razonable.',
        'Diseñar politica salarial competitiva para retener talento especializado.',
        'Elegir provincia segun incentivos, disponibilidad laboral, costos y complejidad tributaria.',
    ]:
        add_bullet(doc, item)
    page_break(doc)

    add_h1(doc, '4. BIBLIOGRAFIA')
    add_p(doc, 'La bibliografia se limita a referencias presentes dentro de los documentos entregados y utilizadas por esos materiales. No se incorporaron nuevas fuentes externas para elaborar el informe.')
    for item in [
        'Ministerio de Economia de la Nacion Argentina. Informe de Cadenas de Valor: Software y Servicios Informaticos N°79, diciembre 2024, citado en los materiales de comercio internacional y SSI.',
        'INDEC. Informes y series sobre Balanza de Pagos, Intercambio Comercial Argentino, IPC, PBI y pobreza citados en los documentos de comercio internacional, PBI y pobreza.',
        'BCRA. Informacion sobre tipo de cambio real, reservas, mercado cambiario y variables financieras citada dentro de los materiales entregados.',
        'CEPAL/CENIA. Indice Latinoamericano de Inteligencia Artificial 2024, citado en los documentos de comercio internacional y tecnologia.',
        'Sysarmy/OpenQube. Encuestas de salarios IT citadas en informe_salarios_argentina y presentacion_slides.',
        'Materiales de catedra de Ingenieria Economica sobre flujo de inversiones, VAN, TIR, ROI, Payback, tasa de descuento, financiamiento y riesgo.',
    ]:
        add_bullet(doc, item)
    page_break(doc)

    add_h1(doc, '5. FUENTES DE DATOS')
    add_p(doc, 'Los siguientes archivos entregados fueron utilizados como base de datos y evidencia para el informe. Se listan para mostrar trazabilidad y separar claramente las fuentes del trabajo de cualquier informacion externa no utilizada.')
    add_table(doc, ['Archivo entregado', 'Uso dentro del informe'], [
        ('Comercio-Internacional-Presentacion (1).pptx', 'Sector externo, balanza de pagos, cuenta corriente, comercio, SSI, EdC, RIGI, talento, infraestructura digital y sintesis estrategica.'),
        ('IE - Comercio Internacional Informe.pdf', 'Desarrollo escrito sobre restriccion externa, comercio SSI, IA, balanza de pagos, regimenes promocionales y condiciones para exportar servicios.'),
        ('Flujo_Inversiones_Presentacion_Economia2026_V4.pptx', 'Herramientas de Ingenieria Economica: flujo de fondos, valor del dinero en el tiempo, VAN, TIR, ROI, Payback, tasa de descuento, financiamiento y riesgo.'),
        ('IE - TPI - PBI.pdf', 'PBI, estructura productiva, comparacion regional, limitaciones del PBI, inversion, informalidad, productividad y analisis PESTEL aplicado al crecimiento.'),
        ('IE- Presentacion - PBI.pdf', 'Evolucion del PBI, stop and go, crecimiento, inversion, restricciones macroeconomicas y lectura institucional.'),
        ('PDF - IE - TP - Pobreza.pdf', 'Pobreza, informalidad, desigualdad, brechas educativas y restricciones sociales para mercado interno y talento futuro.'),
        ('INFORME IE.pdf', 'Instituciones, seguridad juridica, contratos, derecho de competencia, riesgo regulatorio y confianza inversora.'),
        ('PP IE.pptx', 'Rupturas institucionales, libertad economica, marco tributario, RIGI, RIMI y estructura de impuestos.'),
        ('informe_salarios_argentina.pdf', 'Salarios IT, competencia por talento, dolarizacion parcial, seniority y estructura laboral tecnologica.'),
        ('presentacion_slides.pdf', 'Complemento sobre salarios, compensacion IT, productividad marginal, costos laborales y conclusiones socioeconomicas.'),
        ('index.html', 'Material complementario de presentacion/estructura entregado junto con los demas archivos.'),
    ], [2.35, 3.95], caption='Archivos utilizados')

    add_h1(doc, '6. GLOSARIO')
    add_table(doc, ['Termino', 'Definicion usada en el informe'], [
        ('SSI', 'Software y Servicios Informaticos. Sector que desarrolla software, presta servicios digitales y exporta conocimiento tecnologico.'),
        ('Economia del Conocimiento', 'Conjunto de actividades basadas en conocimiento, innovacion, tecnologia, digitalizacion y servicios profesionales de alto valor agregado.'),
        ('Restriccion externa', 'Escasez estructural de divisas que limita crecimiento, importaciones, estabilidad cambiaria e inversion.'),
        ('Cuenta corriente', 'Componente de la balanza de pagos que registra bienes, servicios, rentas y transferencias corrientes.'),
        ('Balanza de pagos', 'Registro de transacciones economicas entre residentes del pais y el resto del mundo.'),
        ('Tipo de cambio real', 'Precio relativo entre bienes locales y externos ajustado por inflacion; incide en competitividad.'),
        ('I+D', 'Investigacion y Desarrollo orientados a generar conocimiento, innovacion, productos o mejoras tecnologicas.'),
        ('Ley de Economia del Conocimiento', 'Regimen argentino de promocion para actividades intensivas en conocimiento con beneficios fiscales y condiciones de acceso.'),
        ('RIGI', 'Regimen de Incentivo para Grandes Inversiones con beneficios fiscales, aduaneros y cambiarios para proyectos incluidos.'),
        ('RIMI', 'Regimen orientado a inversiones de menor escala o PyMEs, mencionado en los materiales como complemento a RIGI.'),
        ('IIBB', 'Ingresos Brutos, impuesto provincial sobre la facturacion.'),
        ('Convenio Multilateral', 'Regimen para distribuir base imponible de Ingresos Brutos entre provincias cuando una empresa opera en varias jurisdicciones.'),
        ('VAN', 'Valor Actual Neto. Indicador que descuenta flujos futuros y evalua si un proyecto genera valor presente positivo.'),
        ('TIR', 'Tasa Interna de Retorno. Tasa que iguala el VAN a cero y estima rentabilidad del proyecto.'),
        ('ROI', 'Retorno sobre la inversion. Relacion entre ganancia obtenida y monto invertido.'),
        ('Payback', 'Plazo de recuperacion de la inversion inicial a partir de los flujos generados por el proyecto.'),
        ('Tasa de descuento', 'Rendimiento minimo exigido por el inversor para traer flujos futuros a valor presente, incorporando riesgo y costo de oportunidad.'),
        ('Nearshoring', 'Estrategia de localizar servicios en paises cercanos por huso horario, afinidad cultural o distancia operativa.'),
        ('IA', 'Inteligencia Artificial. Conjunto de tecnologias capaces de realizar tareas asociadas al aprendizaje, prediccion, lenguaje, percepcion o automatizacion.'),
        ('CTIM/STEM', 'Ciencia, Tecnologia, Ingenieria y Matematica; areas de formacion relevantes para tecnologia y productividad.'),
    ], [1.45, 4.85], caption='Glosario de conceptos principales')

    doc.save(OUT)
    return OUT


if __name__ == '__main__':
    path = build()
    print(path)
