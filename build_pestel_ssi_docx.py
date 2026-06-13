from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/alesso/dev/university/PDC/PESTEL_SSI_Inversion_Argentina.docx')

NAVY = '12355B'
BLUE = '1F6F9B'
LIGHT_BLUE = 'EAF4FB'
LIGHT_GRAY = 'F3F6F8'
GREEN = 'DFF3E8'
RED = 'FBE8E8'
MUTED = '666666'
WHITE = 'FFFFFF'


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, val in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        el = tc_mar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tc_mar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def borders(table, color='D7DEE6'):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = tbl_borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tbl_borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    tr_pr.append(el)


def set_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def white_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(WHITE)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, bold_label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(text[len(bold_label):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table, 'B9CEDD')
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.10
for s in ['List Bullet', 'List Number']:
    styles[s].font.name = 'Calibri'
    styles[s].font.size = Pt(11)
    styles[s].paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('PESTEL Del Estado Actual De Argentina')
r.bold = True
r.font.size = Pt(21)
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Enfoque aplicado a una empresa de software / SSI que evalua invertir')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_callout(doc, 'Recorte del caso', 'El analisis parte del estado actual de Argentina segun los materiales entregados y lo aplica a una empresa internacional o local de software/SSI que analiza invertir, contratar talento y exportar servicios desde el pais. Por eso se priorizan divisas, talento IT, costos salariales, incentivos, infraestructura digital, propiedad intelectual y riesgo regulatorio.', LIGHT_BLUE)

add_h1(doc, 'Estado actual de Argentina que toma el PESTEL')
for item in [
    'Argentina muestra una estabilizacion macroeconomica en proceso, pero todavia condicionada por inflacion, tipo de cambio, riesgo pais, baja profundidad financiera y necesidad de previsibilidad.',
    'El sector externo tiene superavit comercial de bienes en los datos trabajados, pero la restriccion externa sigue presente porque tambien pesan servicios, renta de inversion, reservas y salida de divisas.',
    'El PBI muestra capacidad de recuperacion, aunque con crecimiento heterogeneo por sectores y sin garantia automatica de mejora social si persisten informalidad, pobreza y deterioro salarial.',
    'El mercado laboral esta segmentado: pocos empleos privados formales sostienen gran parte de la estructura, mientras el sector IT paga salarios altos por productividad y competencia global.',
    'El SSI aparece como sector estrategico porque puede transformar talento en exportaciones de servicios, generar divisas y diversificar la matriz productiva fuera de commodities.',
    'El marco institucional combina incentivos concretos, como Ley de Economia del Conocimiento y regimenes provinciales, con riesgos por complejidad tributaria, litigiosidad laboral y cambios regulatorios.',
]:
    add_bullet(doc, item)
add_h1(doc, 'Resumen ejecutivo')
for item in [
    'En el estado actual de Argentina, el atractivo para una empresa SSI no esta principalmente en el mercado interno, sino en usar el pais como plataforma de talento para exportar servicios.',
    'La oportunidad central es generar ingresos en moneda dura con capital humano local, aprovechando salarios competitivos internacionalmente, experiencia exportadora e incentivos a la economia del conocimiento.',
    'El riesgo central es macro-regulatorio: inflacion, tipo de cambio real, normativa cambiaria, presion tributaria, litigiosidad laboral, fuga de talento y cambios de reglas.',
    'La estrategia recomendada es entrar con foco exportador, compensacion competitiva vinculada al dolar, uso de incentivos de Economia del Conocimiento y seleccion cuidadosa de provincia.',
]:
    add_bullet(doc, item)

add_h1(doc, 'Matriz PESTEL del estado actual argentino')
rows = [
    ('Politico', 'Argentina busca ordenar expectativas e impulsar sectores exportadores, pero sigue necesitando previsibilidad institucional.', 'Beneficios fiscales, estabilidad sectorial y apoyo a actividades exportadoras de conocimiento.', 'Volatilidad politica/financiera, riesgo pais, cambios de reglas, controles cambiarios y burocracia.'),
    ('Economico', 'El pais combina recuperacion parcial, inflacion persistente y restriccion externa; necesita divisas genuinas.', 'Exportar SSI genera divisas sin logistica pesada; nearshoring y brecha cambiaria baja favorecen operaciones formales.', 'Apreciacion cambiaria, inflacion salarial, baja profundidad financiera, planificacion incierta y salida de rentas.'),
    ('Social', 'El mercado laboral esta partido entre informalidad amplia y sectores de alta productividad como IT.', 'Capital humano calificado, experiencia exportadora, equipos remotos y salarios competitivos internacionalmente.', 'Fuga de talento, competencia por seniors, informalidad laboral, pobreza infantil y brecha educativa/digital.'),
    ('Tecnologico', 'Existe capacidad tecnologica real, pero concentrada en nichos y limitada por infraestructura, conectividad e I+D.', 'Verticales de alto valor: IA, fintech, salud, AgTech, e-commerce, IoT y productividad empresarial.', 'Rezago 5G, dispositivos caros, baja escala de I+D, falta de capital, ciberseguridad y patentes insuficientes.'),
    ('Ecologico', 'El software tiene baja huella directa, pero IA y data centers vuelven relevante energia, agua y criterios ESG.', 'Software para ESG, energia inteligente, AgTech, ciudades inteligentes e IoT industrial.', 'Disponibilidad electrica, agua, infraestructura verde y licencia social para proyectos intensivos en computo.'),
    ('Legal', 'El marco legal ofrece incentivos a conocimiento e inversion, pero conserva complejidad tributaria, laboral y cambiaria.', 'Exenciones y beneficios tributarios; proteccion de software por propiedad intelectual; modernizacion laboral parcial.', 'Complejidad tributaria multinivel, litigiosidad, cambios cambiarios, demoras administrativas y enforcement legal.'),
]

t = doc.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
headers = ['Factor', 'Estado actual argentino', 'Oportunidades SSI', 'Riesgos SSI']
for i,h in enumerate(headers):
    c=t.rows[0].cells[i]
    c.text=h
    shade(c, BLUE)
    margins(c)
    white_bold(c)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
repeat_header(t.rows[0])
for row in rows:
    cells=t.add_row().cells
    for i,val in enumerate(row):
        cells[i].text=val
        margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i == 0:
            shade(cells[i], LIGHT_GRAY)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor.from_string(NAVY)
borders(t)
set_widths(t, [0.9, 1.75, 1.95, 1.9])

doc.add_paragraph()
add_h1(doc, 'Analisis por factor')

sections = [
    ('1. Factores politicos', [
        ('Estado actual:', 'El pais intenta mejorar previsibilidad macro e incentivar sectores exportadores, pero todavia arrastra volatilidad institucional, riesgo pais y sensibilidad frente a cambios de reglas.'),
        ('Informacion extraida:', 'Los materiales muestran politicas especificas para actividades basadas en conocimiento: Ley de Economia del Conocimiento, RIGI para grandes inversiones y regimenes provinciales. El SSI aparece como via para diversificar exportaciones y aliviar la restriccion externa.'),
        ('Oportunidades:', 'La Ley de Economia del Conocimiento ofrece beneficios fiscales y mayor previsibilidad para software, biotecnologia y servicios intensivos en conocimiento. Algunas provincias, como Cordoba segun los materiales, ofrecen exenciones y programas ligados a empleo CTIM.'),
        ('Riesgos:', 'Los incentivos no eliminan volatilidad politica, riesgo pais, tramites, cambios regulatorios ni incertidumbre cambiaria. Para inversiones SSI, la confianza depende de reglas sostenidas y acceso operativo a divisas.'),
        ('Implicancia para SSI:', 'Conviene estructurar la inversion bajo regimenes promocionales, pero con clausulas de salida, cobertura cambiaria y seguimiento permanente de normativa nacional y provincial.'),
    ]),
    ('2. Factores economicos', [
        ('Estado actual:', 'Argentina muestra recuperacion parcial y superavit comercial de bienes en los datos trabajados, pero mantiene restriccion externa porque la cuenta corriente tambien depende de servicios, renta de inversion, reservas y financiamiento.'),
        ('Informacion extraida:', 'El SSI exporto USD 2.449 millones en 2023 y represento 15,2% de exportaciones de servicios en los materiales de comercio. Para una empresa de software, esto confirma que el sector ya funciona como generador de divisas.'),
        ('Oportunidades:', 'El software genera divisas con baja dependencia de insumos fisicos. La demanda global, el nearshoring, el ingles tecnico y costos relativos pueden convertir Argentina en plataforma exportadora.'),
        ('Riesgos:', 'Inflacion y apreciacion cambiaria reducen margenes; la baja profundidad financiera encarece inversiones; la planificacion con VAN/TIR exige tasas de descuento que reflejen riesgo pais, inflacion y costo de oportunidad.'),
        ('Implicancia para SSI:', 'El modelo mas robusto es export-led: facturacion externa, costos parcialmente locales, compensacion flexible y monitoreo de tipo de cambio real.'),
    ]),
    ('3. Factores sociales', [
        ('Estado actual:', 'La Argentina actual combina pobreza e informalidad elevadas con un segmento tecnologico de alta productividad y salarios por encima del promedio. Esa dualidad define el mercado de talento.'),
        ('Informacion extraida:', 'El sector SSI cuenta con 145.969 empleos registrados en 2023, 6.108 empresas y remuneraciones 78% superiores al promedio privado formal, segun la presentacion de comercio/SSI.'),
        ('Oportunidades:', 'Hay capital humano calificado, polos en CABA, Cordoba, Santa Fe, Mendoza y Rio Negro, y experiencia en trabajo remoto/exportador. Esto permite montar equipos de desarrollo, QA, datos, IA y soporte tecnico.'),
        ('Riesgos:', 'El 32% del sector IT cobra total o parcialmente en dolares; la competencia externa presiona salarios senior. La informalidad superior al 43% y la pobreza infantil afectan el pipeline de talento de largo plazo.'),
        ('Implicancia para SSI:', 'La retencion de talento exige salarios competitivos vinculados al dolar CCL, carrera profesional, beneficios flexibles y formacion interna para ampliar perfiles semi-senior.'),
    ]),
    ('4. Factores tecnologicos', [
        ('Estado actual:', 'El pais posee capacidades tecnologicas reales y talento aplicado, pero el salto de escala depende de infraestructura digital, financiamiento, conectividad y reglas para innovacion.'),
        ('Informacion extraida:', 'Argentina tiene capacidades en software, IA, fintech, salud, AgTech y e-commerce. El SSI invierte 3,6% de ventas en I+D y aporta 14% de la I+D empresaria nacional, segun los materiales.'),
        ('Oportunidades:', 'FONSOFT, FONTAR y FONARSEC pueden apoyar I+D+i. La adopcion de IA y la exportacion de servicios de alto valor permiten competir mas por productividad que por bajo costo.'),
        ('Riesgos:', 'Rezago 5G, dispositivos caros, infraestructura limitada, ciberseguridad y baja generacion de patentes reducen escalabilidad. Startups pueden migrar si falta capital o previsibilidad.'),
        ('Implicancia para SSI:', 'La inversion deberia priorizar verticales de alto valor y arquitectura cloud/remota, evitando depender de infraestructura local critica no garantizada.'),
    ]),
    ('5. Factores ecologicos', [
        ('Estado actual:', 'La agenda ambiental argentina se vincula con energia, agua, agro, mineria y sostenibilidad. Para SSI puro el impacto directo es bajo, pero para IA e infraestructura digital vuelve central la energia.'),
        ('Informacion extraida:', 'El SSI tiene baja huella directa por no depender de logistica pesada ni insumos fisicos intensivos. Pero IA, data centers y computo avanzado elevan demanda electrica y pueden requerir agua para enfriamiento.'),
        ('Oportunidades:', 'Hay demanda global de soluciones ESG: AgTech, energia inteligente, ciudades inteligentes, eficiencia industrial e IoT. Argentina puede combinar software con agro, energia y recursos naturales.'),
        ('Riesgos:', 'Para proyectos intensivos en computo, la disponibilidad electrica, el costo energetico y el agua condicionan localizacion. El PBI no descuenta externalidades ambientales, por lo que la rentabilidad debe revisarse con criterios ESG.'),
        ('Implicancia para SSI:', 'Para software puro el riesgo ambiental es bajo; para IA/data centers debe evaluarse energia, agua, permisos y reputacion ambiental.'),
    ]),
    ('6. Factores legales', [
        ('Estado actual:', 'El marco legal argentino ofrece incentivos para invertir, pero convive con complejidad administrativa, presion tributaria multinivel, litigiosidad laboral y reglas cambiarias sensibles.'),
        ('Informacion extraida:', 'El entorno legal combina incentivos relevantes con complejidad: Ley EdC, RIGI/RIMI, exenciones provinciales, IIBB, Convenio Multilateral, propiedad intelectual, normativa laboral y cambiaria.'),
        ('Oportunidades:', 'Exportaciones de servicios pueden tener tratamiento fiscal favorable en jurisdicciones clave. La propiedad intelectual protege software y la Ley EdC reduce carga fiscal a cambio de I+D, empleo y exportaciones.'),
        ('Riesgos:', 'Persisten multiples niveles tributarios, litigiosidad laboral, cambios de reglas cambiarias, complejidad administrativa por trabajo remoto interprovincial y demoras en registros de PI.'),
        ('Implicancia para SSI:', 'La decision de provincia y estructura societaria es estrategica. Conviene auditoria tributaria-laboral previa y plan de PI para software, marcas, contratos y licencias.'),
    ]),
]

for title, items in sections:
    add_h2(doc, title)
    for label, body in items:
        add_bullet(doc, label + ' ' + body, label)

add_h1(doc, 'Recomendaciones para la empresa SSI')
rec_rows = [
    ('Modelo de negocio', 'Priorizar exportacion de servicios y clientes externos antes que dependencia del mercado interno.'),
    ('Talento', 'Pagar compensaciones competitivas, con componente o referencia dolar CCL, para retener perfiles senior.'),
    ('Localizacion', 'Comparar CABA, Cordoba, Santa Fe, Mendoza y Rio Negro por talento, exenciones, conectividad y costo operativo.'),
    ('Incentivos', 'Inscribirse o evaluar Ley de Economia del Conocimiento; usar RIGI solo si el proyecto tiene escala de infraestructura/gran inversion.'),
    ('Riesgo financiero', 'Evaluar proyectos con VAN/TIR incluyendo inflacion, tipo de cambio, riesgo pais y costo de oportunidad.'),
    ('Tecnologia', 'Focalizar en IA aplicada, fintech, salud, AgTech, e-commerce, datos y automatizacion empresarial.'),
    ('Legal', 'Revisar Convenio Multilateral, IIBB, contratos laborales/remotos, propiedad intelectual y normativa cambiaria.'),
]
rt = doc.add_table(rows=1, cols=2)
rt.alignment = WD_TABLE_ALIGNMENT.CENTER
rt.autofit = False
rt.rows[0].cells[0].text = 'Decision'
rt.rows[0].cells[1].text = 'Recomendacion'
for c in rt.rows[0].cells:
    shade(c, NAVY)
    margins(c)
    white_bold(c)
for a,b in rec_rows:
    c=rt.add_row().cells
    c[0].text=a
    c[1].text=b
    margins(c[0]); margins(c[1])
    shade(c[0], LIGHT_GRAY)
    for p in c[0].paragraphs:
        for r in p.runs:
            r.bold=True
borders(rt)
set_widths(rt, [1.55, 4.95])

add_h1(doc, 'Conclusion ejecutiva')
add_callout(doc, 'Dictamen', 'El estado actual de Argentina presenta una oportunidad selectiva para SSI: hay talento, experiencia exportadora e incentivos, pero tambien inflacion, restriccion externa, complejidad legal y volatilidad regulatoria. La inversion resulta mas defendible si se orienta a exportar servicios, captar divisas, retener talento con compensacion competitiva y usar beneficios de Economia del Conocimiento o regimenes provinciales.', GREEN)

add_h1(doc, 'Archivos considerados')
sources = [
    'Comercio-Internacional-Presentación (1).pptx',
    'IE - Comercio Internacional Informe.pdf',
    'Flujo_Inversiones_Presentacion_Economia2026_V4.pptx',
    'IE - TPI - PBI.pdf',
    'IE-  Presentacion - PBI.pdf',
    'PDF - IE - TP - Pobreza.pdf',
    'index.html',
    'informe_salarios_argentina.pdf',
    'presentacion_slides.pdf',
    'INFORME IE.pdf',
    'PP IE.pptx',
]
for s in sources:
    add_bullet(doc, s)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Nota metodologica: se trabajo solo con los archivos entregados por el usuario y con el recorte especifico de una empresa de software/SSI que evalua invertir en Argentina.')
r.italic = True
r.font.color.rgb = RGBColor.from_string(MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('PESTEL SSI - Ingenieria Economica')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor.from_string(MUTED)

doc.core_properties.title = 'PESTEL Estado Actual Argentina - SSI'
doc.core_properties.subject = 'Empresa de software / SSI'
doc.core_properties.author = 'Codex'
doc.save(OUT)
print(OUT)
