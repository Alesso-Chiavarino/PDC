from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/alesso/dev/university/PDC/FODA_Argentina_Destino_Inversion.docx')

NAVY = '0B2545'
BLUE = '1F4D78'
LIGHT_BLUE = 'EAF3FA'
LIGHT_GREEN = 'E7F3EA'
LIGHT_RED = 'FBE7E7'
LIGHT_GOLD = 'FFF4D6'
LIGHT_GRAY = 'F2F4F7'
WHITE = 'FFFFFF'
BORDER = 'D7DEE6'
MUTED = '555555'


def qn_w(name):
    return qn(f'w:{name}')


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn_w('shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn_w('fill'), fill)


def margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, val in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        node = tc_mar.find(qn_w(side))
        if node is None:
            node = OxmlElement(f'w:{side}')
            tc_mar.append(node)
        node.set(qn_w('w'), str(val))
        node.set(qn_w('type'), 'dxa')


def borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        node = tbl_borders.find(qn_w(edge))
        if node is None:
            node = OxmlElement(f'w:{edge}')
            tbl_borders.append(node)
        node.set(qn_w('val'), 'single')
        node.set(qn_w('sz'), '6')
        node.set(qn_w('space'), '0')
        node.set(qn_w('color'), color)


def set_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement('w:tblHeader')
    node.set(qn_w('val'), 'true')
    tr_pr.append(node)


def setup_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    styles = doc.styles
    styles['Normal'].font.name = 'Calibri'
    styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    styles['Normal'].font.size = Pt(10.8)
    styles['Normal'].paragraph_format.space_after = Pt(7)
    styles['Normal'].paragraph_format.line_spacing = 1.10
    for s in ['List Bullet', 'List Number']:
        styles[s].font.name = 'Calibri'
        styles[s].font.size = Pt(10.8)
        styles[s].paragraph_format.space_after = Pt(4)
        styles[s].paragraph_format.line_spacing = 1.10
    return doc


def style_run(run, bold=False, color=None, size=None, italic=False):
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if size:
        run.font.size = Pt(size)


def title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(72)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Analisis FODA de Argentina como destino de inversion')
    style_run(r, bold=True, color=NAVY, size=22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run('Complemento del informe tecnico PESTEL de Ingenieria Economica')
    style_run(r, italic=True, color=MUTED, size=12.5)

    callout(doc, 'Criterio metodologico', 'El FODA se construye exclusivamente a partir del informe realizado sobre Argentina como destino de inversion y de los materiales trabajados en la materia. No se agregan fuentes externas ni datos nuevos. El enfoque es pais/entorno de inversion para empresas en general; el sector SSI se usa solo como caso testigo cuando ayuda a ejemplificar una capacidad concreta.')


def h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    for r in p.runs:
        style_run(r, bold=True, color=BLUE, size=16)
    return p


def h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        style_run(r, bold=True, color=NAVY, size=13)
    return p


def h3(doc, text):
    p = doc.add_heading(text, level=3)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    for r in p.runs:
        style_run(r, bold=True, color='434343', size=12)
    return p


def para(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    p.add_run(text)
    return p


def callout(doc, heading, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=150, start=190, bottom=150, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(heading)
    style_run(r, bold=True, color=NAVY, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.10
    p2.add_run(body)
    borders(table, 'B7C9D8')
    set_widths(table, [6.3])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, caption=None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(caption)
        style_run(r, bold=True, color=NAVY, size=9.6)
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, head in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = head
        shade(c, BLUE)
        margins(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                style_run(r, bold=True, color=WHITE, size=9.3)
    repeat_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                shade(cells[i], LIGHT_GRAY)
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        style_run(r, bold=True, color=NAVY, size=9.2)
            else:
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.2)
    borders(t)
    set_widths(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def matrix_cell(cell, title_text, items, fill):
    shade(cell, fill)
    margins(cell, top=170, start=190, bottom=170, end=190)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title_text)
    style_run(r, bold=True, color=NAVY, size=12)
    for item in items:
        pp = cell.add_paragraph(style=None)
        pp.paragraph_format.space_after = Pt(3)
        pp.paragraph_format.left_indent = Inches(0.05)
        rr = pp.add_run('- ' + item)
        rr.font.size = Pt(9.4)


def foda_matrix(doc):
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    data = [
        ('FORTALEZAS', [
            'Escala economica y estructura productiva diversificada.',
            'Capacidad exportadora en bienes y servicios.',
            'Talento calificado y sectores de economia del conocimiento.',
            'Regimenes de promocion a la inversion y a la innovacion.',
            'Ecosistema tecnologico con SSI, IA e I+D.'
        ], LIGHT_GREEN),
        ('DEBILIDADES', [
            'Restriccion externa y fragilidad de cuenta corriente.',
            'Inflacion, volatilidad cambiaria y baja profundidad financiera.',
            'Pobreza, informalidad y brechas educativas.',
            'Complejidad tributaria, laboral y administrativa.',
            'Infraestructura digital y financiamiento tecnologico insuficientes.'
        ], LIGHT_RED),
        ('OPORTUNIDADES', [
            'Diversificacion exportadora hacia servicios de alto valor agregado.',
            'Atraccion de inversiones mediante EdC, RIGI, RIMI y provincias.',
            'Crecimiento de IA, fintech, salud, AgTech, e-commerce e IoT.',
            'Desarrollo de polos regionales de talento y empleo formal.',
            'Inversion por etapas en sectores que generen divisas y productividad.'
        ], LIGHT_BLUE),
        ('AMENAZAS', [
            'Cambios regulatorios o institucionales que alteren reglas de juego.',
            'Apreciacion cambiaria que erosione competitividad exportadora.',
            'Fuga de talento y presion salarial en perfiles calificados.',
            'Persistencia de deficit por servicios/rentas y necesidad de financiamiento.',
            'Rezago tecnologico frente a competidores regionales.'
        ], LIGHT_GOLD),
    ]
    idx = 0
    for row in range(2):
        for col in range(2):
            title_text, items, fill = data[idx]
            matrix_cell(t.cell(row, col), title_text, items, fill)
            idx += 1
    borders(t, 'B9C7D3')
    set_widths(t, [3.15, 3.15])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


fortalezas = [
    ('Escala economica y estructura productiva diversificada', 'Argentina no aparece en el informe como una economia pequena ni carente de capacidades. Los materiales sobre PBI y comparacion regional muestran que conserva escala dentro de Sudamerica, con estructura productiva diversificada y presencia de servicios, industria, agro, energia, mineria, software, biotecnologia y servicios profesionales. Esta fortaleza es relevante para empresas porque permite pensar inversiones en distintos sectores y no solo en una actividad aislada.'),
    ('Capacidad exportadora y generacion de divisas', 'El informe destaca que Argentina tiene capacidad de exportar bienes y servicios. El saldo comercial de bienes y los datos de exportaciones muestran una base externa real. Para la inversion, esta fortaleza es clave porque las actividades que generan divisas tienen mejor posicion relativa en un pais con restriccion externa. El SSI funciona como ejemplo: exporto USD 2.449 millones en 2023 y represento 15,2% de las exportaciones de servicios.'),
    ('Capital humano calificado', 'Los materiales muestran una disponibilidad de talento relevante, especialmente en sectores intensivos en conocimiento. En SSI se registran 145.969 empleos y 6.108 empresas, con remuneraciones 78% superiores al promedio privado formal. Esto no debe leerse solo como dato sectorial, sino como senal de que el pais cuenta con capacidades laborales que pueden sostener inversiones en tecnologia, servicios profesionales, administracion, ingenieria y actividades de mayor valor agregado.'),
    ('Regimenes de promocion e incentivos', 'El informe identifica instrumentos como la Ley de Economia del Conocimiento, el RIGI, el RIMI y regimenes provinciales. Estos mecanismos ofrecen beneficios fiscales, estabilidad o mejoras para determinados proyectos. Como fortaleza, no eliminan el riesgo argentino, pero ayudan a reducir costos, mejorar previsibilidad y orientar inversiones hacia actividades estrategicas.'),
    ('Capacidad tecnologica e innovadora', 'Argentina cuenta con capacidades en software, IA, fintech, salud, AgTech, e-commerce e IoT. El sector SSI invierte 3,6% de sus ventas en I+D y aporta 14% de la I+D empresaria nacional. Esta fortaleza muestra que el pais puede atraer inversiones que busquen talento, conocimiento y soluciones digitales para mejorar productividad.')
]

debilidades = [
    ('Restriccion externa y fragilidad de cuenta corriente', 'La principal debilidad macroeconomica es que el superavit comercial de bienes no alcanza por si solo para estabilizar el sector externo. En 1T 2025 la cuenta corriente fue deficitaria en USD 5.191 millones, aun con saldo de bienes positivo. Los deficits de servicios y renta de inversion muestran que el pais necesita financiamiento o generacion adicional de divisas. Para empresas, esto puede afectar importaciones, pagos al exterior, giro de utilidades y disponibilidad de moneda extranjera.'),
    ('Inflacion y volatilidad de precios relativos', 'La inflacion mencionada en los materiales afecta costos, salarios, precios, capital de trabajo y tasa de descuento. En una inversion, esto vuelve mas dificil proyectar flujos de fondos y comparar ingresos futuros con egresos presentes. Desde Ingenieria Economica, esta debilidad impacta directamente sobre VAN, TIR, ROI y Payback.'),
    ('Baja profundidad financiera y costo de capital elevado', 'El informe remarca problemas de credito, financiamiento y baja profundidad financiera. Esto limita la capacidad de las empresas para invertir, ampliar operaciones o financiar capital de trabajo. En proyectos de largo plazo, la falta de credito estable aumenta la dependencia de recursos propios o financiamiento externo, elevando el riesgo.'),
    ('Pobreza, informalidad y brechas educativas', 'La dimension social muestra que el pais tiene talento en nichos, pero tambien pobreza, informalidad y desigualdades educativas. Esta debilidad afecta consumo interno, productividad, formacion futura de trabajadores y formalizacion. Para empresas, puede limitar la demanda local y obligar a invertir mas en capacitacion y retencion.'),
    ('Complejidad tributaria, laboral y administrativa', 'El marco legal argentino combina incentivos con presion tributaria, litigiosidad, Convenio Multilateral, normas laborales y regulaciones cambiarias. Esta debilidad aumenta costos de cumplimiento y exige asesoramiento especializado. No impide invertir, pero reduce la simplicidad operativa y puede alargar plazos de decision.')
]

oportunidades = [
    ('Diversificacion exportadora hacia servicios de alto valor agregado', 'La restriccion externa abre una oportunidad para actividades que generen divisas sin requerir grandes importaciones fisicas. Servicios profesionales, software, economia del conocimiento y exportaciones basadas en talento pueden contribuir a reducir dependencia de commodities. Para empresas, Argentina puede funcionar como plataforma de servicios exportables si se gestiona bien el riesgo cambiario y laboral.'),
    ('Aprovechamiento de incentivos nacionales y provinciales', 'Los regimenes de promocion identificados en el informe permiten mejorar la ecuacion economica de ciertos proyectos. La Ley de Economia del Conocimiento, el RIGI, el RIMI y beneficios provinciales pueden reducir carga fiscal, mejorar estabilidad o favorecer empleo calificado. La oportunidad es seleccionar bien sector y localizacion.'),
    ('Crecimiento tecnologico y adopcion de IA', 'Los materiales destacan IA, fintech, healthtech, AgTech, e-commerce e IoT como areas de potencial. La IA tambien puede aumentar productividad en tareas de oficina y actividades calificadas. Para empresas, esto abre oportunidades de automatizacion, desarrollo de nuevos servicios y mejora de procesos productivos.'),
    ('Desarrollo de polos regionales y talento fuera de CABA', 'Aunque CABA concentra gran parte del empleo SSI, los materiales mencionan crecimiento de Cordoba, Santa Fe, Mendoza y Rio Negro. Esto abre oportunidad para inversiones regionales con menor saturacion, mejores incentivos locales y acceso a talento calificado.'),
    ('Inversion selectiva con ingresos en divisas y Payback razonable', 'El FODA muestra que Argentina es mas atractiva cuando el proyecto genera divisas, tiene capacidad exportadora y recupera la inversion en plazos razonables. Esto permite reducir exposicion a cambios macroeconomicos o regulatorios. La oportunidad no es invertir sin condiciones, sino estructurar proyectos por etapas, con sensibilidad financiera y cobertura de riesgo.')
]

amenazas = [
    ('Volatilidad institucional y cambios regulatorios', 'Los materiales institucionales mencionan antecedentes de alteracion de reglas y contratos. Esta amenaza puede elevar la tasa de descuento exigida por inversores, reducir horizontes de planificacion y postergar proyectos. Afecta especialmente inversiones de largo plazo, intensivas en capital o dependientes de estabilidad normativa.'),
    ('Apreciacion cambiaria y perdida de competitividad', 'El informe senala que el tipo de cambio real afecta competitividad exportadora. Una apreciacion sostenida puede abaratar consumos externos pero erosionar margenes de empresas exportadoras, incluyendo servicios. Esta amenaza es relevante para proyectos que dependen de vender al exterior desde Argentina.'),
    ('Persistencia de deficit por servicios y renta de inversion', 'Aunque exista superavit de bienes, los egresos por servicios y renta pueden mantener fragilidad externa. Esta amenaza puede generar necesidad de financiamiento, presion sobre reservas o restricciones futuras. Para empresas, implica riesgo sobre pagos al exterior, giro de utilidades e importaciones de tecnologia o insumos.'),
    ('Fuga de talento y presion salarial', 'El informe muestra que parte del talento IT cobra total o parcialmente en dolares y que perfiles senior pueden tener brechas relevantes frente a salarios en pesos. Esta amenaza no afecta solo a software: cualquier empresa que necesite perfiles calificados puede enfrentar competencia internacional por talento y aumento de costos de retencion.'),
    ('Rezago tecnologico frente a competidores regionales', 'Argentina tiene capacidades tecnologicas, pero tambien brechas en 5G, conectividad, ciberseguridad, patentes, dispositivos e inversiones entrantes. Si estas brechas persisten, otros paises de la region pueden captar inversiones que requieran infraestructura mas estable o ecosistemas tecnologicos mas maduros.')
]


def develop_list(doc, items):
    for i, (name, body) in enumerate(items, 1):
        h3(doc, f'{i}. {name}')
        para(doc, body)


def build():
    doc = setup_doc()
    title(doc)
    doc.add_page_break()

    h1(doc, '1. Introduccion breve')
    para(doc, 'El presente FODA complementa el informe PESTEL sobre Argentina como destino de inversion. Mientras el PESTEL ordena el entorno en factores politicos, economicos, sociales, tecnologicos y legales, el FODA permite sintetizar ese diagnostico en cuatro categorias estrategicas: fortalezas, debilidades, oportunidades y amenazas.')
    para(doc, 'El objetivo no es elaborar una lista generica, sino interpretar como las caracteristicas actuales del pais afectan a empresas que evaluan invertir. Por eso, cada punto se fundamenta en el informe realizado y en los materiales trabajados: sector externo, PBI, pobreza, salarios, SSI, tecnologia, instituciones, regimenes de promocion y evaluacion de proyectos.')
    para(doc, 'El analisis toma a Argentina como entorno general de inversion. El sector de Software y Servicios Informaticos aparece solo como caso testigo cuando permite mostrar una capacidad concreta del pais: generar divisas, emplear talento calificado, invertir en I+D y participar en economia del conocimiento.')

    h1(doc, '2. Matriz FODA')
    foda_matrix(doc)
    callout(doc, 'Lectura rapida de la matriz', 'El FODA muestra una tension central: Argentina posee capacidades productivas, talento e incentivos, pero esas fortalezas conviven con restricciones macroeconomicas, sociales e institucionales. Por lo tanto, el pais resulta atractivo principalmente para inversiones selectivas, exportadoras, tecnologicas o capaces de elevar productividad.', fill=LIGHT_GOLD)

    h1(doc, '3. Desarrollo analitico')
    h2(doc, '3.1 Fortalezas')
    para(doc, 'Las fortalezas son aspectos internos del pais que pueden mejorar la posicion de Argentina frente a otros destinos de inversion. En este caso, se relacionan con escala economica, diversificacion productiva, capacidad exportadora, talento, incentivos y tecnologia.')
    develop_list(doc, fortalezas)

    h2(doc, '3.2 Debilidades')
    para(doc, 'Las debilidades son limitaciones estructurales internas que reducen el atractivo del pais o encarecen la decision de invertir. No significan que la inversion sea imposible, pero si obligan a exigir mayor rentabilidad, mejor cobertura contractual, escenarios alternativos y plazos de recupero prudentes.')
    develop_list(doc, debilidades)
    doc.add_page_break()

    h2(doc, '3.3 Oportunidades')
    para(doc, 'Las oportunidades surgen de tendencias, sectores dinamicos o condiciones del entorno que pueden ser aprovechadas por empresas. En Argentina, las oportunidades mas claras se vinculan con generacion de divisas, economia del conocimiento, tecnologia, incentivos y desarrollo regional.')
    develop_list(doc, oportunidades)

    h2(doc, '3.4 Amenazas')
    para(doc, 'Las amenazas son riesgos que pueden afectar negativamente la inversion, aun cuando el proyecto tenga fortalezas propias. En el caso argentino, las amenazas se concentran en volatilidad institucional, sector externo, tipo de cambio, talento y rezago tecnologico.')
    develop_list(doc, amenazas)
    doc.add_page_break()

    h1(doc, '4. Integracion del FODA con el informe PESTEL')
    para(doc, 'El FODA no reemplaza al PESTEL, sino que lo resume desde una logica estrategica. Los factores politicos y legales explican buena parte de las fortalezas asociadas a incentivos, pero tambien de las amenazas vinculadas a cambios regulatorios. Los factores economicos explican la fortaleza exportadora y la oportunidad de generar divisas, pero tambien las debilidades de inflacion, restriccion externa y bajo credito.')
    para(doc, 'Los factores sociales aparecen en dos sentidos. Por un lado, el talento calificado es una fortaleza y una oportunidad para atraer servicios basados en conocimiento. Por otro lado, pobreza, informalidad y brechas educativas son debilidades que limitan consumo, productividad y formacion futura. La dimension tecnologica tambien es dual: hay capacidades en IA, software e I+D, pero persisten brechas de infraestructura y adopcion.')
    add_table(doc, ['Relacion PESTEL-FODA', 'Lectura estrategica para inversion'], [
        ('Politico/legal', 'Incentivos como EdC, RIGI, RIMI y provincias fortalecen proyectos, pero la volatilidad regulatoria aumenta riesgo.'),
        ('Economico', 'La escala y exportaciones son fortalezas; la restriccion externa, inflacion y credito limitado son debilidades/amenazas.'),
        ('Social', 'Talento calificado favorece inversion; pobreza, informalidad y fuga de talento limitan escalabilidad.'),
        ('Tecnologico', 'IA, SSI e I+D abren oportunidades; rezago 5G, ciberseguridad y financiamiento tecnologico son restricciones.'),
        ('Evaluacion de proyectos', 'Las variables del FODA deben traducirse a tasa de descuento, sensibilidad de flujos, VAN, TIR, ROI y Payback.'),
    ], [1.7, 4.6], caption='Sintesis de integracion entre PESTEL y FODA')

    h1(doc, '5. Conclusion')
    para(doc, 'El FODA muestra que Argentina es un destino de inversion con potencial real, pero de atractivo selectivo. Sus fortalezas principales son la escala economica, la diversificacion productiva, la capacidad exportadora, el talento calificado, los incentivos legales y las capacidades tecnologicas. Sus debilidades mas importantes son la restriccion externa, la inflacion, la baja profundidad financiera, la informalidad, la pobreza y la complejidad normativa.')
    para(doc, 'Las fortalezas y oportunidades no compensan automaticamente las debilidades y amenazas. Las compensan solo bajo ciertas condiciones: cuando el proyecto genera divisas, incorpora tecnologia, aprovecha talento local, usa regimenes promocionales, controla su exposicion cambiaria y tiene una evaluacion financiera prudente. En cambio, proyectos dependientes del mercado interno, con recupero largo, alto componente importado o baja flexibilidad enfrentan mayor riesgo.')
    para(doc, 'Argentina resulta mas atractiva para invertir cuando la empresa opera en sectores exportadores, economia del conocimiento, servicios profesionales, tecnologia, energia, infraestructura o actividades capaces de aumentar productividad. Tambien mejora su atractivo cuando la inversion se realiza por etapas, con analisis de sensibilidad, tasa de descuento ajustada al riesgo y Payback razonable.')
    callout(doc, 'Respuesta final', 'El FODA confirma la conclusion del informe PESTEL: Argentina no debe analizarse como un destino de bajo riesgo, sino como un pais con capacidades concretas y oportunidades selectivas. La clave para invertir no es ignorar sus debilidades, sino estructurar proyectos que puedan convertir sus fortalezas en flujos de fondos sostenibles.', fill=LIGHT_GREEN)

    h1(doc, '6. Fuentes utilizadas')
    para(doc, 'Este documento utiliza como base el informe tecnico PESTEL de Argentina como destino de inversion y los materiales trabajados en Ingenieria Economica: comercio internacional, flujo de inversiones, PBI, pobreza, salarios, SSI, instituciones, marco tributario y presentaciones complementarias. No se incorporaron fuentes externas nuevas.')
    for item in [
        'Informe Tecnico Final de Ingenieria Economica sobre Argentina como destino de inversion.',
        'Comercio Internacional y sector externo argentino.',
        'IE - Comercio Internacional Informe.',
        'Flujo de Inversiones y Evaluacion de Proyectos.',
        'IE - TPI - PBI e IE - Presentacion - PBI.',
        'PDF - IE - TP - Pobreza.',
        'INFORME IE y PP IE.',
        'informe_salarios_argentina y presentacion_slides.',
    ]:
        bullet(doc, item)

    doc.save(OUT)
    return OUT


if __name__ == '__main__':
    print(build())
