from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/alesso/dev/university/PDC/PESTEL_Argentina_Destino_Inversion.docx')

NAVY = '12355B'
BLUE = '1F6F9B'
LIGHT_BLUE = 'EAF4FB'
LIGHT_GRAY = 'F3F6F8'
GREEN = 'E4F4EA'
WHITE = 'FFFFFF'
MUTED = '666666'


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, val in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        el = tc_mar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tc_mar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def borders(table, color='D7DEE6'):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = tbl_borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tbl_borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    tr_pr.append(el)


def set_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def white_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(WHITE)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet(doc, text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if label and text.startswith(label):
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(text[len(label):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table, 'B9CEDD')
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    return table


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(1)
sec.bottom_margin = Inches(1)
sec.left_margin = Inches(1)
sec.right_margin = Inches(1)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(11)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.10
for s in ['List Bullet', 'List Number']:
    styles[s].font.name = 'Calibri'
    styles[s].font.size = Pt(11)
    styles[s].paragraph_format.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Argentina Como Destino De Inversion')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run('Analisis PESTEL del entorno empresarial con foco aplicado en SSI')
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_callout(doc, 'Nota metodologica', 'El informe se elaboro exclusivamente con los PowerPoint, informes, PDF y documentos entregados. No se incorporan fuentes externas ni datos nuevos. El sector de Software y Servicios Informaticos se usa como ejemplo dinamico dentro de la economia del conocimiento, pero el analisis evalua a Argentina como destino de inversion empresarial en sentido amplio.', LIGHT_BLUE)

add_h1(doc, 'Resumen')
for item in [
    'Argentina presenta un entorno de alta dualidad: recursos productivos, sectores exportadores, capital humano calificado e incentivos a la inversion conviven con inflacion, restriccion externa, volatilidad institucional, pobreza, informalidad y complejidad legal-tributaria.',
    'Para un potencial inversor, las oportunidades aparecen en actividades que generen divisas, incorporen tecnologia, eleven productividad o aprovechen recursos estrategicos: energia, mineria, agroindustria, servicios profesionales, economia del conocimiento y SSI.',
    'El principal aprendizaje de los materiales es que no alcanza con mirar el PBI o un beneficio fiscal aislado. La decision debe integrar riesgo macroeconomico, tipo de cambio, seguridad juridica, financiamiento, disponibilidad de talento, infraestructura y horizonte de recupero.',
    'El SSI funciona como caso testigo: muestra que Argentina puede exportar conocimiento y pagar salarios altos por productividad, pero tambien evidencia fuga de talento, competencia global, necesidad de compensacion dolarizada y dependencia de reglas previsibles.',
]:
    add_bullet(doc, item)

add_h1(doc, 'Introduccion')
intro = [
    'Este informe analiza el estado actual de Argentina como destino de inversion empresarial. El objetivo es ordenar oportunidades y riesgos del entorno de negocios a partir de los materiales de Ingenieria Economica entregados.',
    'El enfoque considera variables politicas, economicas, sociales, tecnologicas y legales. Se incorporan temas de PBI, inflacion, comercio internacional, balanza de pagos, flujo de inversiones, pobreza, salarios, instituciones, propiedad intelectual, infraestructura y economia del conocimiento.',
    'La pregunta guia es: bajo que condiciones Argentina puede ser atractiva para invertir, y que riesgos debe gestionar una empresa antes de radicarse, expandirse o contratar operaciones en el pais?'
]
for txt in intro:
    doc.add_paragraph(txt)

add_h1(doc, 'Analisis PESTEL')

sections = [
    ('1. Factores politicos', [
        ('Informacion extraida:', 'Los materiales muestran una tension entre busqueda de estabilizacion y antecedentes de volatilidad en reglas de juego. Tambien identifican instrumentos publicos de promocion como Ley de Economia del Conocimiento, RIGI y regimenes provinciales.'),
        ('Oportunidades:', 'Los regimenes de promocion pueden reducir costos, mejorar previsibilidad tributaria y favorecer inversiones de largo plazo en tecnologia, infraestructura, energia, mineria o actividades intensivas en conocimiento.'),
        ('Riesgos:', 'Los incentivos formales pueden no compensar por completo riesgo pais, burocracia, cambios regulatorios, controles cambiarios o baja confianza institucional.'),
        ('Impacto empresarial:', 'El factor politico incide sobre tasa de descuento, horizonte de inversion, reinversion de utilidades, eleccion de jurisdiccion y decision de radicacion.'),
        ('Ejemplo SSI:', 'Para software, la Ley de Economia del Conocimiento y los incentivos provinciales ayudan, pero la empresa necesita reglas previsibles para exportar servicios, liquidar divisas, contratar talento y planificar salarios.'),
    ]),
    ('2. Factores economicos', [
        ('Informacion extraida:', 'La economia argentina muestra ciclos de expansion y contraccion asociados a inflacion, crisis cambiarias, baja inversion, restriccion externa y dificultades para sostener credito productivo.'),
        ('Oportunidades:', 'Existe una base productiva diversificada: agroindustria, energia, mineria, litio, industria, servicios profesionales, software, biotecnologia y economia del conocimiento. Las actividades exportadoras son especialmente valiosas porque generan divisas.'),
        ('Riesgos:', 'Inflacion, volatilidad cambiaria, apreciacion del tipo de cambio real, falta de credito, baja profundidad financiera y salida por servicios/rentas elevan la incertidumbre de costos e ingresos.'),
        ('Impacto empresarial:', 'En evaluacion de proyectos, estos riesgos afectan VAN, TIR, ROI, Payback, tasa de descuento, costo de oportunidad y plazo de recupero.'),
        ('Ejemplo SSI:', 'El SSI puede exportar sin logistica pesada y generar divisas, pero su competitividad depende de tipo de cambio real, salarios en dolares, brecha cambiaria, retencion de talento y demanda global.'),
    ]),
    ('3. Factores sociales', [
        ('Informacion extraida:', 'Los documentos muestran una marcada dualidad: sectores de alta productividad y capital humano calificado conviven con pobreza, informalidad, brechas educativas y deterioro del salario real.'),
        ('Oportunidades:', 'Hay disponibilidad de talento calificado en nichos, polos regionales y capacidad de generar empleo formal en actividades dinamicas. Esto favorece servicios, tecnologia, ingenieria, administracion y produccion de mayor valor agregado.'),
        ('Riesgos:', 'Pobreza, informalidad, desigualdad educativa, fuga de talento y competencia por perfiles especializados limitan la ampliacion futura de trabajadores calificados.'),
        ('Impacto empresarial:', 'La dimension social afecta productividad, consumo interno, estabilidad laboral, costos de contratacion, capacitacion y escalabilidad.'),
        ('Ejemplo SSI:', 'Los materiales de salarios IT muestran remuneraciones altas por productividad y competencia global; por eso una empresa SSI debe ofrecer compensacion competitiva, formacion y carrera para retener perfiles senior.'),
    ]),
    ('4. Factores tecnologicos', [
        ('Informacion extraida:', 'Argentina cuenta con sectores dinamicos vinculados a software, servicios profesionales, IA, fintech, salud, AgTech, e-commerce, IoT, biotecnologia y actividades empresariales exportables.'),
        ('Oportunidades:', 'La tecnologia puede elevar productividad, automatizar procesos, mejorar logistica, agregar valor a exportaciones y conectar sectores tradicionales con modelos digitales.'),
        ('Riesgos:', 'Infraestructura insuficiente, rezago 5G, costos de dispositivos, brechas STEM, ciberseguridad, patentes y baja difusion tecnologica pueden frenar la adopcion.'),
        ('Impacto empresarial:', 'La dimension tecnologica define productividad, innovacion, competitividad internacional y capacidad de escalar operaciones modernas.'),
        ('Ejemplo SSI:', 'El SSI aparece como sector capaz de transformar talento en divisas, pero necesita infraestructura digital, I+D, proteccion de propiedad intelectual y financiamiento para startups.'),
    ]),
    ('5. Factores legales', [
        ('Informacion extraida:', 'Los documentos destacan seguridad juridica, estabilidad regulatoria, presion tributaria, normas laborales, restricciones cambiarias, eficacia judicial, propiedad intelectual y regimenes promocionales.'),
        ('Oportunidades:', 'La Ley de Economia del Conocimiento, RIGI, beneficios provinciales, exenciones y marcos de propiedad intelectual pueden mejorar la viabilidad de proyectos con exportaciones, empleo, I+D o infraestructura.'),
        ('Riesgos:', 'Complejidad tributaria, costos laborales, litigiosidad, regulacion cambiante, restricciones cambiarias, Convenio Multilateral y diferencias provinciales aumentan costos de cumplimiento.'),
        ('Impacto empresarial:', 'El marco legal define estructura societaria, contratacion, impuestos, exportaciones, importaciones, financiamiento, distribucion de utilidades y proteccion de activos intangibles.'),
        ('Ejemplo SSI:', 'Para software, son criticos contratos remotos, PI, licencias, tratamiento de exportaciones de servicios, IIBB, normas cambiarias y regimen de Economia del Conocimiento.'),
    ]),
]

for title, items in sections:
    add_h2(doc, title)
    for label, body in items:
        add_bullet(doc, label + ' ' + body, label)

add_h1(doc, 'Matriz integrada de impacto para la inversion')
rows = [
    ('Politico', 'Regimenes de promocion, estabilidad fiscal sectorial, incentivos nacionales y provinciales.', 'Volatilidad institucional, cambios regulatorios, tramites, baja confianza contractual.', 'Alto: determina horizonte de inversion, tasa de descuento y radicacion.'),
    ('Economico', 'Sectores diversificados, recursos naturales, exportaciones, economia del conocimiento y divisas.', 'Inflacion, restriccion externa, volatilidad cambiaria, bajo credito e incertidumbre de costos.', 'Alto: afecta rentabilidad, financiamiento, demanda e importaciones.'),
    ('Social', 'Capital humano calificado en nichos, empleo formal en sectores dinamicos y polos regionales.', 'Pobreza, informalidad, brechas educativas y presion salarial en talento especializado.', 'Medio/Alto: condiciona consumo, productividad y contratacion.'),
    ('Tecnologico', 'IA, I+D, digitalizacion, AgTech, fintech, salud, e-commerce e IoT.', 'Rezago 5G, costos de dispositivos, brechas STEM, ciberseguridad y baja difusion.', 'Alto: define productividad, innovacion y competitividad.'),
    ('Legal', 'Ley de Economia del Conocimiento, RIGI, exenciones y beneficios provinciales.', 'Presion tributaria, costos laborales, restricciones cambiarias y complejidad administrativa.', 'Alto: define costos, formalizacion, estructura operativa y previsibilidad.'),
]

t = doc.add_table(rows=1, cols=4)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
headers = ['Factor', 'Oportunidades', 'Riesgos', 'Impacto sobre la inversion']
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    shade(c, BLUE)
    margins(c)
    white_bold(c)
    c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
repeat_header(t.rows[0])
for row in rows:
    cells = t.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val
        margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i == 0:
            shade(cells[i], LIGHT_GRAY)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor.from_string(NAVY)
borders(t)
set_widths(t, [0.9, 1.9, 1.9, 1.8])

doc.add_paragraph()
add_h1(doc, 'Conclusiones')
for item in [
    'Argentina resulta atractiva para invertir de manera selectiva, no automatica. La oportunidad depende del sector, la localizacion, el acceso a divisas, el marco legal y la capacidad de operar en volatilidad.',
    'Las principales ventajas son recursos productivos, capital humano, sectores exportadores, potencial energetico-minero, agroindustria, servicios profesionales, economia del conocimiento y regimenes de promocion.',
    'Los principales riesgos son inflacion, restriccion externa, volatilidad cambiaria, baja inversion, presion tributaria, informalidad, pobreza, complejidad legal e incertidumbre regulatoria.',
    'Para mejorar su atractivo, Argentina deberia consolidar estabilidad macroeconomica, fortalecer seguridad juridica, ampliar infraestructura fisica y digital, mejorar credito, sostener incentivos productivos y formar capital humano.',
    'El SSI confirma la tesis general: Argentina tiene capacidades reales para competir globalmente, pero el inversor debe disenar una estrategia prudente con gestion de riesgo cambiario, laboral, tributario y regulatorio.',
]:
    add_bullet(doc, item)

add_h1(doc, 'Fuentes de datos utilizadas')
sources = [
    ('Comercio-Internacional-Presentacio\u0301n (1).pptx', 'Restriccion externa, cuenta corriente, condiciones para exportar SSI, politica sectorial, infraestructura digital, energia/ESG y sintesis estrategica.'),
    ('IE - Comercio Internacional Informe.pdf', 'Balanza de pagos, comercio internacional, participacion argentina en SSI, exportacion de servicios y entorno para SSI.'),
    ('Flujo_Inversiones_Presentacion_Economia2026_V4.pptx', 'Flujo de fondos, valor del dinero en el tiempo, VAN, TIR, ROI, Payback, riesgo y tasa de descuento.'),
    ('IE - TPI - PBI.pdf', 'PBI, inflacion, inversion, informalidad, tecnologia, marco legal e institucional.'),
    ('IE-  Presentacion - PBI.pdf', 'Evolucion del PBI, stop and go, inflacion, PESTEL aplicado al PBI y dimension legal-institucional.'),
    ('PDF - IE - TP - Pobreza.pdf', 'Pobreza, informalidad, brechas educativas, dualidad social y limites para la formacion futura de talento.'),
    ('index.html', 'Presentacion web sobre pobreza, trampas estructurales, capital humano, tecnologia e inflacion.'),
    ('informe_salarios_argentina.pdf', 'Salarios IT, competencia por talento, dolarizacion parcial, productividad marginal y mercado laboral tecnologico.'),
    ('presentacion_slides.pdf', 'Metricas de compensacion IT, seniority, brecha cambiaria, productividad y conclusiones economico-sociales.'),
    ('INFORME IE.pdf', 'Instituciones, seguridad juridica, volatilidad regulatoria, contratos, propiedad intelectual y riesgo institucional.'),
    ('PP IE.pptx', 'Propiedad intelectual, economia del conocimiento, innovacion, patentes, IA y derecho de autor.'),
]
st = doc.add_table(rows=1, cols=2)
st.alignment = WD_TABLE_ALIGNMENT.CENTER
st.autofit = False
st.rows[0].cells[0].text = 'Archivo'
st.rows[0].cells[1].text = 'Uso dentro del informe'
for c in st.rows[0].cells:
    shade(c, NAVY)
    margins(c)
    white_bold(c)
for a, b in sources:
    c = st.add_row().cells
    c[0].text = a
    c[1].text = b
    margins(c[0]); margins(c[1])
    shade(c[0], LIGHT_GRAY)
    for p in c[0].paragraphs:
        for r in p.runs:
            r.bold = True
borders(st)
set_widths(st, [2.35, 4.15])

doc.add_paragraph()
add_h1(doc, 'Glosario')
gloss = [
    ('SSI', 'Software y Servicios Informaticos. Sector que desarrolla software, presta servicios digitales y exporta conocimiento tecnologico.'),
    ('Economia del Conocimiento', 'Actividades basadas en conocimiento, digitalizacion, innovacion, tecnologia y servicios profesionales de alto valor agregado.'),
    ('Restriccion externa', 'Escasez estructural de divisas que limita crecimiento, importaciones, estabilidad cambiaria e inversion.'),
    ('Cuenta corriente', 'Componente de la balanza de pagos que registra bienes, servicios, rentas y transferencias corrientes.'),
    ('Tipo de cambio real', 'Precio relativo entre bienes locales y externos, ajustado por inflacion; incide en competitividad.'),
    ('Nearshoring', 'Estrategia de instalar o contratar operaciones en paises cercanos en huso horario, cultura o distancia operativa.'),
    ('I+D', 'Investigacion y Desarrollo orientados a generar conocimiento, innovacion, productos, procesos o mejoras tecnologicas.'),
    ('VAN', 'Valor Actual Neto. Indicador que descuenta flujos futuros y evalua si un proyecto genera valor presente positivo.'),
    ('TIR', 'Tasa Interna de Retorno. Tasa que iguala el VAN a cero y estima la rentabilidad promedio del proyecto.'),
    ('ROI', 'Retorno sobre la inversion. Relacion entre ganancia obtenida y monto invertido.'),
    ('Payback', 'Plazo de recuperacion de la inversion inicial a partir de los flujos generados por el proyecto.'),
]
gt = doc.add_table(rows=1, cols=2)
gt.alignment = WD_TABLE_ALIGNMENT.CENTER
gt.autofit = False
gt.rows[0].cells[0].text = 'Termino'
gt.rows[0].cells[1].text = 'Definicion'
for c in gt.rows[0].cells:
    shade(c, BLUE)
    margins(c)
    white_bold(c)
for a, b in gloss:
    c = gt.add_row().cells
    c[0].text = a
    c[1].text = b
    margins(c[0]); margins(c[1])
    shade(c[0], LIGHT_GRAY)
    for p in c[0].paragraphs:
        for r in p.runs:
            r.bold = True
borders(gt)
set_widths(gt, [1.8, 4.7])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Nota final: el informe evita citas de pagina no verificadas y usa nombres de archivos reales entregados para mantener trazabilidad sin agregar fuentes externas.')
r.italic = True
r.font.color.rgb = RGBColor.from_string(MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('Argentina como destino de inversion - PESTEL')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor.from_string(MUTED)

doc.core_properties.title = 'Argentina como destino de inversion - PESTEL'
doc.core_properties.subject = 'Entorno empresarial con foco SSI'
doc.core_properties.author = 'Codex'
doc.save(OUT)
print(OUT)
