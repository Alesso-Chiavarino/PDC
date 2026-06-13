from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path('/Users/alesso/dev/university/PDC/PESTEL_Argentina_Destino_Inversion.docx')

NAVY = '12355B'
BLUE = '1F6F9B'
LIGHT_BLUE = 'EAF4FB'
LIGHT_GRAY = 'F3F6F8'
GREEN = 'E4F4EA'
WHITE = 'FFFFFF'
MUTED = '666666'


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for side, val in {'top': top, 'start': start, 'bottom': bottom, 'end': end}.items():
        el = tc_mar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tc_mar.append(el)
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')


def borders(table, color='D7DEE6'):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.first_child_found_in('w:tblBorders')
    if tbl_borders is None:
        tbl_borders = OxmlElement('w:tblBorders')
        tbl_pr.append(tbl_borders)
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = tbl_borders.find(qn(f'w:{edge}'))
        if el is None:
            el = OxmlElement(f'w:{edge}')
            tbl_borders.append(el)
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '6')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    el = OxmlElement('w:tblHeader')
    el.set(qn('w:val'), 'true')
    tr_pr.append(el)


def set_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Inches(w)


def white_bold(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(WHITE)


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor.from_string(NAVY)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_p(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, label=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if label and text.startswith(label):
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(text[len(label):])
    else:
        p.add_run(text)
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    borders(table, 'B9CEDD')
    cell = table.cell(0, 0)
    shade(cell, fill)
    margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        shade(c, BLUE)
        margins(c)
        white_bold(c)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    repeat_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i == 0:
                shade(cells[i], LIGHT_GRAY)
                for p in cells[i].paragraphs:
                    for r in p.runs:
                        r.bold = True
                        r.font.color.rgb = RGBColor.from_string(NAVY)
    borders(t)
    set_widths(t, widths)
    return t


def page(doc, title=None):
    if title:
        add_h1(doc, title)


def next_page(doc):
    doc.add_page_break()


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.9)
sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
styles['Normal'].font.size = Pt(10.8)
styles['Normal'].paragraph_format.space_after = Pt(6)
styles['Normal'].paragraph_format.line_spacing = 1.08
for s in ['List Bullet', 'List Number']:
    styles[s].font.name = 'Calibri'
    styles[s].font.size = Pt(10.8)
    styles[s].paragraph_format.space_after = Pt(4)

# PAGE 1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(120)
p.paragraph_format.space_after = Pt(8)
r = p.add_run('Argentina Como Destino De Inversion')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor.from_string(NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(18)
r = p.add_run('Analisis PESTEL del entorno empresarial con foco aplicado en SSI')
r.italic = True
r.font.size = Pt(13)
r.font.color.rgb = RGBColor.from_string(MUTED)

add_callout(doc, 'Nota metodologica', 'El informe se elaboro exclusivamente con los PowerPoint, informes, PDF y documentos entregados. No se incorporan fuentes externas ni datos nuevos. Por indicacion del usuario, se excluye el apartado ecologico/ambiental y se desarrolla el analisis en dimensiones politica, economica, social, tecnologica y legal.')
for item in [
    'Objeto: evaluar Argentina como destino de inversion empresarial.',
    'Caso testigo: Software y Servicios Informaticos (SSI) como sector dinamico de la economia del conocimiento.',
    'Criterio: ordenar oportunidades, riesgos e impactos sobre decisiones de inversion.',
]:
    add_bullet(doc, item)
next_page(doc)

# PAGE 2
page(doc, 'Resumen ejecutivo')
for item in [
    'Argentina presenta un entorno de alta dualidad. Por un lado, posee recursos productivos, sectores exportadores, capacidades tecnologicas, capital humano calificado y regimenes de promocion. Por otro lado, enfrenta inflacion, restriccion externa, volatilidad institucional, pobreza, informalidad, baja profundidad financiera y complejidad legal-tributaria.',
    'La inversion empresarial en Argentina no debe evaluarse solo por tamano de mercado o potencial sectorial. La decision requiere analizar tipo de cambio real, acceso a divisas, estabilidad normativa, estructura laboral, financiamiento, infraestructura, impuestos y velocidad de recupero del capital.',
    'El PBI y el crecimiento agregado son indicadores utiles, pero los materiales remarcan sus limites: pueden ocultar desigualdad, informalidad, deterioro salarial, diferencias sectoriales y falta de bienestar. Por eso el PESTEL debe complementar la mirada macro con condiciones institucionales y sociales.',
    'El SSI se usa como ejemplo de oportunidad concreta. Es un sector capaz de exportar conocimiento, generar divisas, pagar salarios superiores al promedio y operar con baja dependencia de insumos fisicos. Sin embargo, tambien expone riesgos de fuga de talento, competencia internacional por perfiles senior, necesidad de compensaciones dolarizadas y dependencia de reglas cambiarias previsibles.',
]:
    add_bullet(doc, item)
add_callout(doc, 'Sintesis', 'Argentina puede ser atractiva para invertir, pero de manera selectiva. Los proyectos mas defendibles son aquellos con capacidad exportadora, tecnologia, productividad, cobertura frente a inflacion y reglas legales/tributarias revisadas antes de escalar.', GREEN)
next_page(doc)

# PAGE 3
page(doc, 'Introduccion y alcance')
for txt in [
    'El presente informe analiza el estado actual de Argentina como destino de inversion empresarial. La lectura integra informacion de los materiales sobre comercio internacional, PBI, pobreza, salarios IT, flujo de inversiones, instituciones, propiedad intelectual y economia del conocimiento.',
    'La pregunta principal no es si Argentina es buena o mala para invertir en terminos absolutos. La pregunta correcta es bajo que condiciones, en que sectores, con que estructura juridica y con que estrategia de riesgo una empresa puede operar de manera rentable.',
    'El informe mantiene un enfoque amplio: sirve para pensar inversiones en sectores productivos, servicios, tecnologia, energia, mineria, agroindustria y actividades profesionales. Dentro de ese marco, el SSI aparece como caso testigo porque los documentos lo presentan como un sector con potencial exportador, empleo calificado y rol en la generacion de divisas.',
    'Se excluye el apartado ecologico/ambiental por indicacion del usuario. Por ese motivo, los impactos vinculados a energia o infraestructura solo se mencionan cuando afectan directamente inversion, tecnologia, costos o localizacion, pero no se desarrolla una dimension ambiental independiente.'
]:
    add_p(doc, txt)
add_h2(doc, 'Estructura del informe')
for item in ['Contexto actual argentino.', 'Analisis politico, economico, social, tecnologico y legal.', 'Matriz integrada de impacto.', 'Aplicacion al caso SSI.', 'Criterios para evaluar proyectos.', 'Recomendaciones, conclusiones, fuentes y glosario.']:
    add_bullet(doc, item)
next_page(doc)

# PAGE 4
page(doc, 'Contexto actual de Argentina')
for item in [
    'Macroeconomia: los materiales muestran una economia marcada por ciclos de expansion y contraccion, inflacion persistente, crisis cambiarias y restricciones de financiamiento. La recuperacion del PBI puede existir, pero no necesariamente implica mejora homogenea ni sostenida.',
    'Sector externo: el superavit comercial de bienes no garantiza equilibrio externo si existen salidas por servicios, renta de inversion, turismo, pagos externos o reduccion de reservas. Esta lectura es central para entender por que las actividades generadoras de divisas tienen valor estrategico.',
    'Mercado laboral: la economia combina empleo formal limitado, informalidad elevada y sectores especificos de alta productividad. En tecnologia, los salarios se explican por productividad marginal, competencia global y posibilidad de exportar servicios.',
    'Instituciones: los documentos sobre entorno legal e institucional remarcan que seguridad juridica, contratos, propiedad intelectual, impuestos y regulacion laboral afectan la confianza de largo plazo.',
    'Inversion: los materiales de flujo de inversiones muestran que los proyectos deben analizarse con VAN, TIR, ROI, Payback, riesgo, tasa de descuento y valor del dinero en el tiempo. En Argentina, esos calculos deben incorporar inflacion, tipo de cambio, riesgo pais y escenarios alternativos.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':')
next_page(doc)

# PAGE 5
page(doc, 'Indicadores cuantitativos incorporados')
add_p(doc, 'Para reforzar el informe, esta seccion agrega datos numericos extraidos de las presentaciones e informes entregados. No se agregan fuentes externas: los valores provienen del material de comercio internacional, PBI, pobreza, salarios IT, flujo de inversiones, presentacion_slides e informes institucionales.')
add_table(doc, ['Bloque', 'Indicadores cuantitativos usados', 'Lectura para inversion'], [
    ('Sector externo', 'Cuenta corriente -USD 5.191 MM; bienes +USD 2.060 MM; servicios -USD 4.502 MM; renta -USD 3.333 MM; cuenta financiera +USD 7.229 MM.', 'El superavit de bienes no alcanza por si solo. La inversion exportadora es mas atractiva si genera divisas genuinas y reduce exposicion a servicios/rentas.'),
    ('PBI e inflacion', 'PBI real 2025 +4,4%; PBI nominal 2024-2025 de 583.909 a 847.622 millones de pesos corrientes; deflactor implicito 39,1%.', 'La recuperacion debe leerse con cautela: una parte del aumento nominal responde a precios y no a mayor produccion real.'),
    ('Comercio 2026', 'Exportaciones marzo 2026 USD 8.645 MM; importaciones USD 6.122 MM; saldo USD 2.523 MM; superavit 1T 2026 USD 5.508 MM.', 'Hay capacidad exportadora, pero la sostenibilidad depende de precios, cantidades, tipo de cambio y composicion de la cuenta corriente.'),
    ('SSI', 'Exportaciones SSI USD 2.449 MM en 2023; 15,2% de exportaciones de servicios; 145.969 empleos registrados; 6.108 empresas; salarios +78% vs promedio privado formal.', 'El SSI tiene escala real y valor estrategico: genera divisas, empleo calificado y productividad superior.'),
    ('Mercado laboral IT', '32% del sector cobra total/parcialmente en USD; brecha CCL/oficial 3,5%; mediana IT $3.253.581; senior $3.500.000; junior $1.538.500.', 'Retener talento exige politica salarial competitiva y flexible. El costo laboral IT debe medirse contra ingresos exportables.'),
    ('Pobreza y capital humano', 'Pobreza oficial 28,2%; indigencia 6,3%; pobreza infantil urbana 53,6%; inseguridad alimentaria infantil 28,8%; 16% de ninos pobres con internet de calidad.', 'La condicion social limita mercado interno y pipeline futuro de talento, aunque existan nichos altamente productivos.'),
], [1.25, 2.55, 2.7])
next_page(doc)

# PAGE 6
page(doc, 'Datos cuantitativos: sector externo y comercio')
add_p(doc, 'Los documentos de comercio internacional muestran que la fragilidad externa argentina no se explica solo por exportaciones e importaciones de bienes. La cuenta corriente integra bienes, servicios, rentas y transferencias; por eso un saldo comercial positivo puede convivir con deficit externo.')
add_table(doc, ['Indicador', 'Valor en los materiales', 'Interpretacion'], [
    ('Cuenta corriente 1T 2025', '-USD 5.191 MM', 'Deficit externo pese a saldo positivo de bienes.'),
    ('Saldo de bienes 1T 2025', '+USD 2.060 MM', 'Superavit comercial insuficiente para cubrir otros egresos.'),
    ('Servicios 1T 2025', '-USD 4.502 MM', 'Principal salida asociada, entre otros factores, a viajes y servicios.'),
    ('Renta de inversion', '-USD 3.333 MM', 'Remision de rentas/utilidades presiona la cuenta corriente.'),
    ('Cuenta financiera', '+USD 7.229 MM', 'El equilibrio depende de financiamiento externo.'),
    ('Exportaciones marzo 2026', 'USD 8.645 MM', 'Record exportador mencionado en los materiales.'),
    ('Importaciones marzo 2026', 'USD 6.122 MM', 'Crecimiento moderado, con cantidades en caida segun el material.'),
    ('Saldo marzo 2026', '+USD 2.523 MM', 'Resultado comercial positivo de corto plazo.'),
    ('Superavit 1T 2026', '+USD 5.508 MM', 'Buen dato comercial, pero no garantiza equilibrio externo total.'),
], [1.8, 1.55, 3.15])
add_callout(doc, 'Lectura cuantitativa', 'Los numeros muestran por que una inversion que genere divisas tiene valor estrategico. Para Argentina, no alcanza con producir mas: importa producir bienes o servicios capaces de sostener la cuenta externa.', GREEN)
next_page(doc)

# PAGE 7
page(doc, 'Datos cuantitativos: SSI y tecnologia')
add_p(doc, 'El sector SSI aparece en los materiales como ejemplo de actividad capaz de transformar talento en exportaciones. Los datos cuantitativos permiten justificar por que se lo usa como caso testigo dentro del analisis de inversion.')
add_table(doc, ['Indicador SSI / tecnologia', 'Valor', 'Implicancia'], [
    ('Exportaciones SSI 2023', 'USD 2.449 MM', 'Genera divisas sin depender de logistica fisica pesada.'),
    ('Participacion en servicios', '15,2%', 'Peso relevante dentro de exportaciones de servicios.'),
    ('Saldo SSI', '+USD 107 MM', 'Sector con contribucion externa positiva segun materiales.'),
    ('Maximo exportador reciente', 'USD 2.609 MM en 2022', 'Muestra capacidad de escala previa.'),
    ('Participacion mundial Argentina', '0,59% en 2011; 0,27%/0,3% en 2023', 'Hay potencial, pero tambien perdida relativa de posicion global.'),
    ('Empleo registrado SSI', '145.969 empleos en 2023', 'Base laboral formal y especializada.'),
    ('Empresas SSI', '6.108 empresas', 'Ecosistema empresarial existente.'),
    ('Remuneraciones SSI', '+78% sobre promedio privado formal', 'Sector de alta productividad relativa.'),
    ('I+D SSI', '3,6% de ventas; 14% de I+D empresaria nacional', 'Capacidad innovadora superior al promedio.'),
    ('ILIA 2024', 'Argentina 47,4; Chile 64,6; Uruguay 60,7; Brasil 52,5', 'Argentina tiene capacidades, pero rezaga frente a lideres regionales.'),
], [2.0, 1.65, 2.85])
next_page(doc)

# PAGE 8
page(doc, 'Datos cuantitativos: salarios, pobreza e inversion')
add_p(doc, 'Los datos laborales y sociales ayudan a evitar una lectura excesivamente optimista. Argentina tiene talento calificado, pero tambien restricciones sociales que afectan consumo, estabilidad laboral y formacion futura de capital humano.')
add_table(doc, ['Tema', 'Dato cuantitativo', 'Lectura'], [
    ('Mercado formal', '6,2 millones de asalariados privados formales sobre 46 millones de habitantes', 'Base formal reducida respecto de la poblacion total.'),
    ('Informalidad', 'Mas del 43% de ocupados en informalidad general segun materiales de salarios', 'Limita productividad, recaudacion y proteccion social.'),
    ('Dolarizacion IT', '32% cobra total/parcialmente en divisas', 'Competencia global por talento y necesidad de compensacion flexible.'),
    ('Brecha cambiaria', 'CCL $1.480 vs oficial $1.430; brecha 3,5%', 'Reduce distorsion para ingreso formal de divisas en el escenario presentado.'),
    ('Salario senior IT', '$3.500.000; USD 2.365 al CCL', 'Costo alto localmente, pero potencialmente competitivo para exportar.'),
    ('Salario mediano IT', '$3.253.581; USD 2.198 al CCL', 'Referencia para presupuesto de equipos tecnologicos.'),
    ('Perdida real IT ejemplo', 'Ajuste nominal 13% vs inflacion 19% = -5,04% real', 'La inflacion exige politicas salariales dinamicas.'),
    ('Pobreza general', '28,2%; indigencia 6,3%', 'El mercado interno y bienestar siguen condicionados.'),
    ('Pobreza infantil', '53,6%; indigencia infantil 10,7%', 'Riesgo para capital humano futuro.'),
    ('Evaluacion financiera', 'VAN, TIR, ROI y Payback como criterios de decision', 'Los proyectos deben cuantificar riesgo, tiempo y retorno esperado.'),
], [1.65, 2.25, 2.6])
next_page(doc)

# PAGE 9
page(doc, 'Lectura de los indicadores cuantitativos')
for item in [
    'Los datos de comercio exterior muestran una oportunidad clara para inversiones exportadoras, pero tambien revelan que la restriccion externa es mas amplia que la balanza comercial de bienes.',
    'Los indicadores SSI muestran que el sector no es solo una promesa: ya tiene exportaciones, empleo formal, empresas, salarios por encima del promedio y participacion relevante en I+D.',
    'Los datos salariales muestran que invertir en tecnologia exige presupuestar talento en terminos competitivos internacionalmente, no solo en pesos nominales.',
    'Los datos de pobreza e informalidad advierten que el crecimiento puede no traducirse automaticamente en bienestar ni en ampliacion rapida del capital humano disponible.',
    'Los criterios de inversion financiera obligan a convertir el analisis PESTEL en supuestos numericos: tasa de descuento, inflacion esperada, tipo de cambio, salarios, impuestos, ventas, recupero y escenarios de estres.',
]:
    add_bullet(doc, item)
add_callout(doc, 'Conclusion de datos', 'Los numeros refuerzan la conclusion: Argentina no es un destino simple, pero si puede ser atractivo para proyectos con productividad, exportacion, gestion de riesgo e incentivos bien aprovechados.', GREEN)
next_page(doc)

# PAGE 10
page(doc, '1. Factores politicos - Estado actual')
for item in [
    'Informacion extraida: los documentos describen un pais que busca estabilizacion y promocion de sectores estrategicos, pero que carga con antecedentes de cambios normativos y debilidad institucional relativa frente a economias regionales mas previsibles.',
    'La politica economica influye directamente sobre expectativas. Cuando el Estado logra ordenar cuentas, sostener reglas y reducir incertidumbre, mejora la tasa de descuento de los proyectos. Cuando predomina la volatilidad, los inversores exigen retornos mayores o reducen plazos.',
    'Los materiales sobre comercio internacional y SSI destacan politicas como Ley de Economia del Conocimiento, RIGI para grandes inversiones y regimenes provinciales. Estas herramientas buscan atraer capital, empleo, exportaciones e innovacion.',
    'La promocion sectorial es relevante, pero no reemplaza estabilidad macro e institucional. Una empresa puede recibir beneficios, pero igual quedar expuesta a cambios cambiarios, demoras administrativas, riesgo pais o incertidumbre sobre repatriacion de utilidades.',
]:
    add_bullet(doc, item, item.split(':')[0] + ':' if ':' in item else None)
add_table(doc, ['Elemento politico', 'Efecto sobre inversion'], [
    ('Estabilidad fiscal', 'Reduce incertidumbre y facilita proyectos de largo plazo.'),
    ('Riesgo pais', 'Aumenta tasa de descuento, costo de financiamiento y retorno exigido.'),
    ('Promocion sectorial', 'Mejora atractivo en tecnologia, energia, infraestructura y conocimiento.'),
    ('Coordinacion Nacion-provincias', 'Afecta impuestos, permisos, beneficios y costos de cumplimiento.'),
], [2.1, 4.4])
next_page(doc)

# PAGE 6
page(doc, '1. Factores politicos - Oportunidades y riesgos')
add_h2(doc, 'Oportunidades')
for item in [
    'Regimenes de promocion que reducen costos o aportan previsibilidad tributaria.',
    'Politicas publicas orientadas a economia del conocimiento, exportaciones de servicios y sectores con capacidad de generar divisas.',
    'Competencia entre provincias para atraer empresas mediante exenciones, programas de empleo y beneficios locales.',
    'Posibilidad de estructurar proyectos con beneficios si se cumplen requisitos de empleo, inversion, exportacion o I+D.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Riesgos')
for item in [
    'Cambios de reglas que afecten planificacion, contratos, importaciones, giro de dividendos o regimen cambiario.',
    'Burocracia y demoras que elevan costos de entrada.',
    'Riesgo pais y volatilidad financiera que encarecen capital.',
    'Desconfianza inversora si los incentivos no se perciben sostenibles.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Implicancia para SSI')
add_p(doc, 'Una empresa de software o servicios informaticos puede aprovechar beneficios publicos, pero debe analizar si esos incentivos son suficientes frente al riesgo macro-politico. La decision deberia incluir revision de regimenes aplicables, jurisdiccion provincial, condiciones de exportacion, contratos laborales y mecanismos de cobertura.')
next_page(doc)

# PAGE 7
page(doc, '2. Factores economicos - Estado actual')
for item in [
    'Informacion extraida: los materiales de PBI muestran una economia con capacidad productiva, pero afectada por inflacion, baja inversion, stop and go, restriccion externa y volatilidad cambiaria.',
    'El PBI mide produccion agregada, pero no alcanza para medir bienestar ni calidad de crecimiento. La recuperacion puede estar concentrada en algunos sectores y no trasladarse automaticamente a salarios reales, empleo formal o reduccion de pobreza.',
    'El sector externo es un condicionante central. Los materiales de comercio internacional muestran que un saldo positivo de bienes no necesariamente resuelve la cuenta corriente si existen salidas por servicios, renta de inversion o perdida de reservas.',
    'Para la inversion, el problema economico no es solo vender. Tambien importa importar insumos, acceder a credito, pagar salarios, girar utilidades, protegerse de inflacion y sostener margenes frente al tipo de cambio real.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':' if ':' in item else None)
add_table(doc, ['Variable economica', 'Por que importa'], [
    ('Inflacion', 'Distorsiona costos, precios, salarios y flujos futuros.'),
    ('Tipo de cambio real', 'Define competitividad exportadora y costo relativo local.'),
    ('Credito', 'Condiciona expansion, capital de trabajo e inversion.'),
    ('Cuenta corriente', 'Mide fragilidad externa mas alla del comercio de bienes.'),
    ('PBI real', 'Ayuda a medir actividad, pero debe complementarse con empleo e ingresos.'),
], [2.0, 4.5])
next_page(doc)

# PAGE 8
page(doc, '2. Factores economicos - Evaluacion de proyectos')
add_p(doc, 'Los materiales de flujo de inversiones son clave porque permiten traducir el entorno argentino a una decision financiera. En un pais volatil, los proyectos no pueden analizarse solo por ingresos esperados; deben incorporar riesgo, tiempo y costo de oportunidad.')
for item in [
    'VAN: permite descontar flujos futuros. En Argentina, la tasa de descuento debe reflejar inflacion, riesgo pais, incertidumbre cambiaria y costo de capital.',
    'TIR: permite comparar rentabilidad estimada con el rendimiento minimo requerido. Si el entorno exige retornos altos, proyectos que parecen atractivos pueden dejar de serlo.',
    'ROI: mide relacion entre resultado e inversion. Es util, pero incompleto si no considera tiempos, riesgo y volatilidad.',
    'Payback: resulta especialmente importante en contextos inciertos porque mide cuanto tarda en recuperarse el capital. En Argentina, plazos largos requieren mayor previsibilidad.',
    'Escenarios: conviene estimar escenario base, optimista y estresado, modificando inflacion, tipo de cambio, salarios, demanda, impuestos y acceso a divisas.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':')
add_callout(doc, 'Lectura para inversores', 'La inversion mas robusta es aquella que genera divisas, reduce dependencia de credito local, recupera capital en plazos razonables y puede ajustar precios o contratos frente a inflacion.', GREEN)
next_page(doc)

# PAGE 9
page(doc, '2. Factores economicos - Oportunidades y riesgos')
add_h2(doc, 'Oportunidades')
for item in [
    'Base productiva diversificada: agroindustria, energia, mineria, litio, industria, servicios profesionales, software y biotecnologia.',
    'Sectores exportadores capaces de generar divisas y aliviar restriccion externa.',
    'Economia del conocimiento con potencial de vender servicios al exterior sin grandes importaciones fisicas.',
    'Costos relativos que pueden ser competitivos para empresas internacionales si se administra bien el riesgo cambiario.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Riesgos')
for item in [
    'Inflacion que reduce poder adquisitivo y dificulta planificacion.',
    'Apreciacion cambiaria que puede erosionar competitividad exportadora.',
    'Baja profundidad financiera y credito limitado para escalar proyectos.',
    'Restriccion externa que puede afectar importaciones, pagos, servicios y utilidades.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Implicancia para SSI')
add_p(doc, 'El SSI es atractivo porque puede facturar en moneda dura, usar talento local y operar con baja logistica. Pero sus costos principales son salarios y retencion de talento, por lo que la estrategia debe cuidar margenes frente a inflacion y tipo de cambio.')
next_page(doc)

# PAGE 10
page(doc, '3. Factores sociales - Estado actual')
for item in [
    'Informacion extraida: los documentos sobre pobreza, PBI y salarios muestran una sociedad dual. Existen segmentos de alta productividad y capital humano calificado, pero tambien pobreza, informalidad, brechas educativas y fragilidad del salario real.',
    'La pobreza no afecta solo al consumo inmediato. Tambien limita capital humano futuro, salud, educacion, conectividad y capacidad de incorporarse a sectores dinamicos.',
    'La informalidad reduce productividad, proteccion social, recaudacion y acceso al credito. Tambien dificulta que el crecimiento se transforme en bienestar generalizado.',
    'El sector IT muestra una realidad distinta: salarios superiores al promedio, competencia global y valoracion de experiencia tecnica. Esto confirma que Argentina tiene nichos de talento, pero no una base social homogenea.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':' if ':' in item else None)
add_table(doc, ['Dimension social', 'Impacto empresarial'], [
    ('Pobreza', 'Limita consumo interno y formacion futura de talento.'),
    ('Informalidad', 'Reduce productividad y formalizacion de proveedores/trabajadores.'),
    ('Capital humano', 'Permite actividades de mayor valor agregado en nichos.'),
    ('Salario real', 'Afecta demanda, clima laboral y negociacion salarial.'),
    ('Brecha educativa', 'Condiciona disponibilidad futura de perfiles tecnicos.'),
], [2.0, 4.5])
next_page(doc)

# PAGE 11
page(doc, '3. Factores sociales - Talento y mercado laboral')
add_p(doc, 'Para una empresa, la dimension social no es solo una variable de responsabilidad social. Es una condicion productiva. La disponibilidad de trabajadores formados, su costo, su estabilidad y su capacidad de aprender tecnologias determinan escalabilidad.')
for item in [
    'En sectores tradicionales, la informalidad y los costos laborales pueden reducir productividad y dificultar cumplimiento.',
    'En sectores dinamicos, como SSI, la competencia por talento puede elevar remuneraciones y hacer necesaria una estrategia de retencion.',
    'La formacion continua es clave porque la brecha entre perfiles junior, semi-senior y senior define costos y tiempos de entrega.',
    'El trabajo remoto externo genera oportunidades para profesionales argentinos, pero tambien presiona a empresas locales a pagar en condiciones comparables al mercado global.',
    'Las universidades y polos regionales son activos relevantes para localizar operaciones, especialmente si se combinan con incentivos provinciales.'
]:
    add_bullet(doc, item)
add_callout(doc, 'Lectura social', 'Argentina tiene talento, pero no sobra talento listo para escalar cualquier operacion. La inversion debe incluir capacitacion, retencion, carrera profesional y politicas salariales adaptadas a inflacion y dolarizacion parcial.', GREEN)
next_page(doc)

# PAGE 12
page(doc, '3. Factores sociales - Oportunidades y riesgos')
add_h2(doc, 'Oportunidades')
for item in [
    'Capital humano calificado en tecnologia, servicios profesionales, ingenieria y administracion.',
    'Polos regionales que permiten diversificar localizacion y costos.',
    'Posibilidad de generar empleo formal en sectores de alta productividad.',
    'Experiencia de trabajo remoto y exportacion de servicios.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Riesgos')
for item in [
    'Fuga de talento hacia empleadores externos.',
    'Necesidad de compensaciones dolarizadas o indexadas para perfiles criticos.',
    'Limitaciones educativas y digitales que reducen pipeline futuro.',
    'Desigualdad territorial que concentra oportunidades en pocos centros urbanos.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Implicancia para SSI')
add_p(doc, 'En software, el costo central es talento. Por eso la empresa que invierte en Argentina debe pensar mas alla del salario nominal: seniority, dolar CCL, beneficios, flexibilidad, capacitacion, cultura, retencion y competencia global.')
next_page(doc)

# PAGE 13
page(doc, '4. Factores tecnologicos - Estado actual')
for item in [
    'Informacion extraida: Argentina posee capacidades en software, IA, fintech, salud, agroindustria, e-commerce, IoT, biotecnologia y servicios profesionales. Sin embargo, enfrenta limites en infraestructura digital, conectividad, 5G, patentes y financiamiento.',
    'La tecnologia aparece como oportunidad transversal porque permite aumentar productividad, exportar servicios, automatizar procesos, mejorar logistica y agregar valor a sectores tradicionales.',
    'Los materiales de SSI remarcan inversion en I+D y rol del sector dentro de la I+D empresaria nacional. Esto muestra que el conocimiento puede ser un factor exportador y no solo un soporte interno.',
    'El desarrollo tecnologico no se distribuye automaticamente. Puede concentrarse en nichos de alta productividad mientras otras empresas quedan rezagadas por falta de capital, conectividad o capacidades.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':' if ':' in item else None)
add_table(doc, ['Activo tecnologico', 'Potencial de inversion'], [
    ('Software', 'Exportacion de servicios, automatizacion y productividad.'),
    ('IA', 'Nuevos productos, analitica, eficiencia y servicios avanzados.'),
    ('Fintech', 'Soluciones financieras en economia con baja bancarizacion relativa.'),
    ('AgTech', 'Valor agregado al agro y mejora de eficiencia productiva.'),
    ('I+D', 'Diferenciacion y creacion de propiedad intelectual.'),
], [2.0, 4.5])
next_page(doc)

# PAGE 14
page(doc, '4. Factores tecnologicos - Infraestructura e innovacion')
add_p(doc, 'Una inversion moderna depende de infraestructura digital. La presencia de talento no alcanza si faltan conectividad, seguridad, equipamiento, financiamiento y reglas de propiedad intelectual.')
for item in [
    'Conectividad: el rezago en 5G o redes avanzadas puede limitar IA en tiempo real, IoT, robotica industrial y operaciones intensivas en datos.',
    'Ciberseguridad: a mayor digitalizacion, mayor necesidad de proteger datos, continuidad operativa y confianza de clientes internacionales.',
    'Propiedad intelectual: el software, las marcas, las licencias y los desarrollos tecnologicos requieren proteccion legal clara.',
    'Financiamiento: startups y empresas tecnologicas necesitan capital para escalar. Si falta financiamiento local, pueden migrar o vender su propiedad intelectual afuera.',
    'Adopcion: la tecnologia mejora productividad solo si las empresas tradicionales la incorporan. Si queda limitada a nichos, el impacto macro es menor.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':')
add_callout(doc, 'Lectura tecnologica', 'Argentina tiene creatividad y talento, pero el salto de escala exige infraestructura, financiamiento, seguridad juridica de intangibles y mayor difusion tecnologica al entramado productivo.', GREEN)
next_page(doc)

# PAGE 15
page(doc, '4. Factores tecnologicos - Oportunidades y riesgos')
add_h2(doc, 'Oportunidades')
for item in [
    'Desarrollo de soluciones para IA, fintech, salud, agro, comercio electronico y servicios empresariales.',
    'Exportacion de software y servicios profesionales con baja necesidad de insumos importados.',
    'Automatizacion de procesos productivos en empresas locales.',
    'Creacion de propiedad intelectual local con potencial internacional.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Riesgos')
for item in [
    'Infraestructura digital insuficiente para ciertos proyectos intensivos.',
    'Brechas de habilidades STEM y alfabetizacion tecnologica.',
    'Ciberseguridad y proteccion de datos como costos crecientes.',
    'Patentes, registros y enforcement lentos o complejos.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Implicancia para SSI')
add_p(doc, 'El SSI es una de las mejores vias para convertir talento en exportaciones. Pero para sostener competitividad necesita infraestructura digital, actualizacion permanente, proteccion de PI, acceso a clientes externos y equipos preparados para estandares internacionales.')
next_page(doc)

# PAGE 16
page(doc, '5. Factores legales - Estado actual')
for item in [
    'Informacion extraida: el marco legal argentino combina incentivos relevantes con alta complejidad. Los materiales mencionan seguridad juridica, estabilidad regulatoria, presion tributaria, normas laborales, restricciones cambiarias, propiedad intelectual y eficacia judicial.',
    'La existencia de beneficios como Ley de Economia del Conocimiento, RIGI o regimenes provinciales puede mejorar la viabilidad de proyectos, pero requiere cumplir condiciones y administrar riesgos de cambios normativos.',
    'La carga tributaria no depende solo de impuestos nacionales. Tambien intervienen provincias y municipios, por lo que la localizacion puede cambiar el costo real de operar.',
    'El regimen laboral incide sobre formalizacion, contratacion, litigiosidad, costo de despido, contribuciones patronales y flexibilidad operativa.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':' if ':' in item else None)
add_table(doc, ['Aspecto legal', 'Impacto empresarial'], [
    ('Seguridad juridica', 'Define confianza y horizonte de largo plazo.'),
    ('Impuestos', 'Afectan costo operativo y rentabilidad neta.'),
    ('Normas laborales', 'Condicionan contratacion, despido y litigiosidad.'),
    ('Reglas cambiarias', 'Afectan exportaciones, pagos y giro de utilidades.'),
    ('Propiedad intelectual', 'Protege software, marcas, patentes y licencias.'),
], [2.0, 4.5])
next_page(doc)

# PAGE 17
page(doc, '5. Factores legales - Oportunidades y riesgos')
add_h2(doc, 'Oportunidades')
for item in [
    'Regimenes promocionales que reducen carga fiscal o dan estabilidad.',
    'Tratamiento favorable para actividades de conocimiento y exportaciones de servicios en ciertas jurisdicciones.',
    'Marco de propiedad intelectual aplicable a software, obras, marcas y desarrollos.',
    'Modernizacion parcial de instrumentos laborales y tributarios en los materiales analizados.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Riesgos')
for item in [
    'Complejidad tributaria multinivel y diferencias provinciales.',
    'Litigiosidad laboral y costos de formalizacion.',
    'Cambios en normativa cambiaria o fiscal.',
    'Demoras administrativas en registros, permisos o enforcement de derechos.'
]:
    add_bullet(doc, item)
add_h2(doc, 'Implicancia para SSI')
add_p(doc, 'En SSI, lo legal es especialmente importante porque el principal activo es intangible. Contratos de desarrollo, licencias, propiedad intelectual, confidencialidad, datos, exportacion de servicios, IIBB, Convenio Multilateral y modalidad laboral/remota deben definirse antes de escalar.')
next_page(doc)

# PAGE 18
page(doc, 'Matriz integrada de impacto para la inversion')
add_table(doc, ['Factor', 'Oportunidades', 'Riesgos', 'Impacto'], [
    ('Politico', 'Regimenes de promocion, estabilidad fiscal sectorial, incentivos nacionales y provinciales.', 'Volatilidad institucional, cambios regulatorios, tramites y baja confianza contractual.', 'Alto: determina horizonte, tasa de descuento y radicacion.'),
    ('Economico', 'Sectores diversificados, recursos naturales, exportaciones, economia del conocimiento y divisas.', 'Inflacion, restriccion externa, volatilidad cambiaria, bajo credito e incertidumbre de costos.', 'Alto: afecta rentabilidad, financiamiento, demanda e importaciones.'),
    ('Social', 'Capital humano calificado en nichos, empleo formal en sectores dinamicos y polos regionales.', 'Pobreza, informalidad, brechas educativas y presion salarial en talento especializado.', 'Medio/Alto: condiciona consumo, productividad y contratacion.'),
    ('Tecnologico', 'IA, I+D, digitalizacion, AgTech, fintech, salud, e-commerce e IoT.', 'Rezago 5G, costos de dispositivos, brechas STEM, ciberseguridad y baja difusion.', 'Alto: define productividad, innovacion y competitividad.'),
    ('Legal', 'Ley de Economia del Conocimiento, RIGI, exenciones y beneficios provinciales.', 'Presion tributaria, costos laborales, restricciones cambiarias y complejidad administrativa.', 'Alto: define costos, formalizacion, estructura operativa y previsibilidad.'),
], [0.9, 1.9, 1.9, 1.8])
add_h2(doc, 'Lectura de la matriz')
add_p(doc, 'La inversion empresarial en Argentina se vuelve mas atractiva cuando combina generacion de divisas, productividad, uso de talento local y aprovechamiento de incentivos. Se vuelve mas riesgosa cuando depende solo del mercado interno, tiene plazo largo de recupero, requiere importaciones criticas o queda expuesta a cambios regulatorios.')
next_page(doc)

# PAGE 19
page(doc, 'Sectores con potencial de inversion segun los materiales')
add_p(doc, 'Una mejora clara del documento comparado es explicitar que Argentina no ofrece una unica oportunidad sectorial. Los materiales muestran una base productiva diversificada: agroindustria, mineria, energia, servicios, industria, software, biotecnologia y actividades basadas en conocimiento. Esta seccion ordena esos sectores sin abrir un apartado ecologico independiente.')
add_table(doc, ['Sector', 'Por que aparece como oportunidad', 'Condiciones o riesgos a gestionar'], [
    ('Agroindustria', 'Aporta divisas, encadenamientos con transporte, industria alimenticia, comercio exterior y productividad exportadora.', 'Dependencia de precios internacionales, clima, logistica, financiamiento e infraestructura.'),
    ('Energia e hidrocarburos', 'Los materiales mencionan recursos no convencionales como Vaca Muerta y potencial para aumentar exportaciones e inversion.', 'Necesita reglas de largo plazo, infraestructura, financiamiento y estabilidad para proyectos intensivos en capital.'),
    ('Mineria y litio', 'Puede atraer inversion, generar exportaciones y diversificar la matriz productiva.', 'Requiere seguridad juridica, permisos, infraestructura y horizonte de recupero largo.'),
    ('Servicios profesionales', 'El predominio del sector servicios y la disponibilidad de capital humano permiten exportar conocimiento y tareas de mayor valor agregado.', 'Depende de talento, idioma, conectividad, tipo de cambio real y competencia regional.'),
    ('Software / SSI', 'Genera divisas con baja necesidad de insumos fisicos, empleo calificado y salarios superiores al promedio.', 'Fuga de talento, salarios dolarizados, infraestructura digital, regulacion cambiaria y competencia internacional.'),
    ('Biotecnologia y conocimiento', 'Aparece vinculada a la Ley de Economia del Conocimiento y a actividades intensivas en I+D.', 'Necesita financiamiento, propiedad intelectual, capital humano especializado y estabilidad regulatoria.'),
    ('Infraestructura y logistica', 'Es necesaria para que el crecimiento sectorial se transforme en productividad y exportaciones sostenidas.', 'Proyectos de plazo largo expuestos a tasa de descuento, riesgo pais, permisos y costos de capital.'),
], [1.55, 2.5, 2.45])
add_callout(doc, 'Lectura sectorial', 'La inversion mas atractiva no surge de un promedio nacional, sino de seleccionar sectores con capacidad de generar divisas, productividad y valor agregado. SSI es un caso central, pero no el unico: energia, mineria, agroindustria y servicios tambien aparecen en los materiales como motores potenciales.', GREEN)
next_page(doc)

# PAGE 20
page(doc, 'Aplicacion: SSI como caso testigo')
add_p(doc, 'El sector SSI permite observar la mejor cara y los principales riesgos del clima de inversion argentino. No es el unico sector relevante, pero funciona como ejemplo porque conecta talento, exportaciones, tecnologia, salarios, propiedad intelectual e incentivos publicos.')
for item in [
    'Fortaleza: puede transformar capital humano en divisas sin depender de logistica pesada.',
    'Fortaleza: existe experiencia exportadora y demanda global de servicios informaticos.',
    'Fortaleza: los salarios IT son altos localmente, pero todavia pueden ser competitivos frente a mercados desarrollados.',
    'Riesgo: la competencia internacional por talento obliga a compensaciones dolarizadas o indexadas.',
    'Riesgo: la apreciacion cambiaria puede reducir margen exportador.',
    'Riesgo: cambios cambiarios, tributarios o laborales pueden alterar el modelo de negocio.',
    'Condicion: la inversion SSI debe orientarse a clientes externos, propiedad intelectual protegida, contratos claros, retencion de talento y uso de incentivos disponibles.'
]:
    add_bullet(doc, item, item.split(':')[0] + ':')
add_callout(doc, 'Conclusion del caso SSI', 'Argentina no es un destino de bajo riesgo, pero si puede ser un destino competitivo para operaciones SSI exportadoras, especialmente si la empresa gestiona riesgo cambiario, retencion de talento y cumplimiento legal.', GREEN)
next_page(doc)

# PAGE 20
page(doc, 'Criterios de decision para un inversor')
add_p(doc, 'Antes de invertir, una empresa deberia transformar el PESTEL en preguntas operativas. Esto evita que el analisis quede como descripcion general y lo convierte en herramienta de decision.')
add_table(doc, ['Area de decision', 'Preguntas clave'], [
    ('Mercado', 'El proyecto depende del mercado interno o puede exportar? La demanda esta dolarizada o pesificada?'),
    ('Finanzas', 'El VAN sigue siendo positivo con inflacion alta, devaluacion, suba salarial o caida de demanda?'),
    ('Divisas', 'Como se cobran exportaciones? Como se pagan proveedores externos? Se pueden girar utilidades?'),
    ('Talento', 'Hay perfiles disponibles? Como se retienen seniors? Que parte del salario se referencia al dolar?'),
    ('Legal', 'Que regimen promocional aplica? Como impactan IIBB, Convenio Multilateral, contratos y PI?'),
    ('Localizacion', 'Que provincia ofrece mejor combinacion de talento, incentivos, conectividad y costos?'),
], [1.65, 4.85])
next_page(doc)

# PAGE 21
page(doc, 'Recomendaciones estrategicas')
for item in [
    'Priorizar proyectos con capacidad exportadora o ingresos vinculados a moneda dura.',
    'Usar escenarios financieros conservadores: base, optimista y estresado.',
    'Aplicar VAN, TIR, ROI y Payback con tasas de descuento que reflejen riesgo argentino.',
    'Elegir localizacion segun talento, beneficios fiscales, conectividad y costos de cumplimiento.',
    'Revisar estructura tributaria nacional, provincial y municipal antes de invertir.',
    'Proteger propiedad intelectual, contratos, datos y licencias desde el inicio.',
    'Diseñar politica de compensacion flexible para perfiles criticos.',
    'Evitar depender exclusivamente de subsidios o incentivos que puedan cambiar.',
    'Invertir en capacitacion para ampliar base de talento y reducir dependencia de seniors escasos.',
    'Mantener estrategia de cobertura frente a inflacion, tipo de cambio y cambios normativos.'
]:
    add_bullet(doc, item)
next_page(doc)

# PAGE 22
page(doc, 'Conclusiones')
for item in [
    'Argentina resulta atractiva para invertir de manera selectiva, no generalizada ni automatica.',
    'Su atractivo se basa en sectores con potencial exportador, recursos productivos, talento calificado, capacidades tecnologicas, economia del conocimiento y regimenes de promocion.',
    'Los riesgos principales son inflacion, restriccion externa, volatilidad cambiaria, baja inversion, presion tributaria, informalidad, pobreza, complejidad legal e incertidumbre regulatoria.',
    'El pais ofrece oportunidades importantes, pero exige evaluacion rigurosa de sector, localizacion, estructura legal, exposicion cambiaria, horizonte de recupero y capacidad de operar en volatilidad.',
    'El SSI confirma que Argentina puede competir globalmente cuando combina conocimiento, productividad y exportacion de servicios. Al mismo tiempo, muestra que retener talento, sostener competitividad y navegar reglas cambiarias son condiciones centrales.',
    'En sintesis, Argentina puede ser buen destino de inversion si el proyecto tiene estrategia clara, gestion de riesgo y capacidad de generar valor en sectores dinamicos.'
]:
    add_bullet(doc, item)
next_page(doc)

# PAGE 23
page(doc, 'Fuentes de datos utilizadas')
add_table(doc, ['Archivo', 'Uso dentro del informe'], [
    ('Comercio-Internacional-Presentacion (1).pptx', 'Restriccion externa, cuenta corriente, condiciones para exportar SSI, politica sectorial, infraestructura digital y sintesis estrategica.'),
    ('IE - Comercio Internacional Informe.pdf', 'Balanza de pagos, comercio internacional, participacion argentina en SSI, exportacion de servicios y entorno para SSI.'),
    ('Flujo_Inversiones_Presentacion_Economia2026_V4.pptx', 'Flujo de fondos, valor del dinero en el tiempo, VAN, TIR, ROI, Payback, riesgo y tasa de descuento.'),
    ('IE - TPI - PBI.pdf', 'PBI, inflacion, inversion, informalidad, tecnologia y marco legal-institucional.'),
    ('IE- Presentacion - PBI.pdf', 'Evolucion del PBI, stop and go, inflacion, PESTEL aplicado al PBI y dimension legal-institucional.'),
    ('PDF - IE - TP - Pobreza.pdf', 'Pobreza, informalidad, brechas educativas, dualidad social y limites para la formacion futura de talento.'),
    ('index.html', 'Presentacion web sobre pobreza, trampas estructurales, capital humano, tecnologia e inflacion.'),
    ('informe_salarios_argentina.pdf', 'Salarios IT, competencia por talento, dolarizacion parcial, productividad marginal y mercado laboral tecnologico.'),
    ('presentacion_slides.pdf', 'Metricas de compensacion IT, seniority, brecha cambiaria, productividad y conclusiones economico-sociales.'),
    ('INFORME IE.pdf', 'Instituciones, seguridad juridica, volatilidad regulatoria, contratos, propiedad intelectual y riesgo institucional.'),
    ('PP IE.pptx', 'Propiedad intelectual, economia del conocimiento, innovacion, patentes, IA y derecho de autor.'),
], [2.35, 4.15])
next_page(doc)

# PAGE 24
page(doc, 'Glosario')
add_table(doc, ['Termino', 'Definicion'], [
    ('SSI', 'Software y Servicios Informaticos. Sector que desarrolla software, presta servicios digitales y exporta conocimiento tecnologico.'),
    ('Economia del Conocimiento', 'Actividades basadas en conocimiento, digitalizacion, innovacion, tecnologia y servicios profesionales de alto valor agregado.'),
    ('Restriccion externa', 'Escasez estructural de divisas que limita crecimiento, importaciones, estabilidad cambiaria e inversion.'),
    ('Cuenta corriente', 'Componente de la balanza de pagos que registra bienes, servicios, rentas y transferencias corrientes.'),
    ('Tipo de cambio real', 'Precio relativo entre bienes locales y externos, ajustado por inflacion; incide en competitividad.'),
    ('Nearshoring', 'Estrategia de instalar o contratar operaciones en paises cercanos en huso horario, cultura o distancia operativa.'),
    ('I+D', 'Investigacion y Desarrollo orientados a generar conocimiento, innovacion, productos, procesos o mejoras tecnologicas.'),
    ('VAN', 'Valor Actual Neto. Indicador que descuenta flujos futuros y evalua si un proyecto genera valor presente positivo.'),
    ('TIR', 'Tasa Interna de Retorno. Tasa que iguala el VAN a cero y estima la rentabilidad promedio del proyecto.'),
    ('ROI', 'Retorno sobre la inversion. Relacion entre ganancia obtenida y monto invertido.'),
    ('Payback', 'Plazo de recuperacion de la inversion inicial a partir de los flujos generados por el proyecto.'),
], [1.8, 4.7])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
r = p.add_run('Nota final: el documento mantiene la exclusion del apartado ecologico/ambiental solicitada por el usuario y utiliza solo los materiales entregados.')
r.italic = True
r.font.color.rgb = RGBColor.from_string(MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run('Argentina como destino de inversion - PESTEL adaptado')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor.from_string(MUTED)

doc.core_properties.title = 'Argentina como destino de inversion - PESTEL adaptado'
doc.core_properties.subject = 'Entorno empresarial con foco SSI, sin apartado ecologico'
doc.core_properties.author = 'Codex'
doc.save(OUT)
print(OUT)
